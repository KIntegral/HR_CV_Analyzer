import pymupdf
from PIL import Image
import pytesseract
import ollama
import json
import os
import re

from docx.oxml import parse_xml
from fpdf import FPDF

from PIL import ImageEnhance, ImageFilter, Image
import numpy as np
from io import BytesIO
from reportlab.platypus import Paragraph, Spacer

from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

from reportlab.lib.pagesizes import A4

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Dla Windows odkomentuj:
pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'

logo_path = r"C:\Users\Kamil Czyżewski\OneDrive - Integral Solutions sp. z o.o\Pulpit\Projects\HR_CV_Analyzer\IS_New.png" #"/app/IS_New.png" #

from ollama import Client
ollama_host = os.getenv('OLLAMA_HOST', 'http://localhost:11434')
ollama_client = Client(host=ollama_host)

try:
    # Próba użycia systemowych fontów
    pdfmetrics.registerFont(TTFont('DejaVuSans', 'DejaVuSans.ttf'))
    pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', 'DejaVuSans-Bold.ttf'))
except:
    # Jeśli nie ma DejaVu, użyj domyślnych z ReportLab
    from reportlab.lib.fonts import addMapping
    # ReportLab ma wbudowane fonty z podstawowym wsparciem UTF-8


def try_fill_company_period_from_text(section, job):
    """
    Uzupełnia company/period na podstawie surowego tekstu sekcji/całego CV.
    Obsługuje układ:
        <job title>
        <firmy | ... | daty>
    oraz:
        <job title> - <firma>, <daty>
    """
    if job.get("company") and job.get("period"):
        return job

    position = (job.get("position") or "").strip()
    if not position:
        return job

    lines = section.splitlines()

    for idx, line in enumerate(lines):
        if position.lower() in line.lower():
            candidate_lines = [line]
            if idx + 1 < len(lines):
                candidate_lines.append(lines[idx + 1])

            combined = " ".join(candidate_lines)

            # okres
            m_period = re.search(
                r"(\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)"
                r"\s+\d{4}\s*[-–]\s*(?:current|present|"
                r"(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}))",
                combined,
                re.IGNORECASE,
            )
            if m_period and not job.get("period"):
                job["period"] = m_period.group(1).strip()

            # firma = wszystko przed okresem (albo cała linia, jeśli dat brak)
            if m_period:
                before_period = combined[:m_period.start()].strip()
            else:
                before_period = combined.strip()

            if before_period and not job.get("company"):
                # wytnij sam tytuł stanowiska
                low = before_period.lower()
                pos_low = position.lower()
                if pos_low in low:
                    start = low.find(pos_low)
                    before_period = before_period[start + len(position):].strip(" -–|,")
                if len(before_period) > 2:
                    job["company"] = before_period

            break

    return job


def safe_text(text, default="NA"):
    if text is None or text == "":
        return default
    if isinstance(text, list):  # ✅ OBSŁUGA LIST
        return ", ".join([str(x) for x in text if x])
    return str(text).strip()

class CVAnalyzer:
    def __init__(self, model_name="qwen2.5:14b"):
        self.model_name = model_name
        
        self.TECH_KNOWLEDGE_BASE = {
            "programming": [
                "Python", "Java", "C#", "JavaScript", "TypeScript", "C++", "Go", "Ruby", "PHP", 
                "Swift", "Kotlin", "Rust", "C", "Scala", "R", "MATLAB", "Julia", "Bash", "PowerShell",
                "Perl", "Haskell", "Erlang", "Elixir", "Clojure", "Groovy", "VB.NET", "F#",
                "Objective-C", "Dart", "Lua", "Coffeescript", "Assembly", "VBA"
            ],
            "frameworks": [
                "Django", "Flask", "FastAPI", "Spring", "Spring Boot", "Express", "NestJS", "Fastify",
                "React", "Angular", "Vue", "Svelte", "Next.js", "Remix", "SvelteKit",
                "PyTorch", "TensorFlow", "Keras", "JAX", "Scikit-learn", "Pandas", "NumPy",
                "Polars", "H2O.ai", "XGBoost", "LightGBM", "CatBoost",
                "Rails", "Laravel", "Symfony", "Django REST", "GraphQL",
                "Pytest", "Jest", "Mocha", "Jasmine", "JUnit", "TestNG", "Mockito",
                "Selenium", "Cypress", "Playwright", "RxJava", "Coroutines", "Compose", "Jetpack",
                "Hibernate", "JPA", "Sequelize", "TypeORM", "Prisma", "SQLAlchemy",
                "Maven", "Gradle", "npm", "yarn", "pip", "conda", "cargo",
                "MVVM", "MVP", "MVI", "CLEAN", "Hexagonal"
            ],
            "mobile": [
                "Android", "iOS", "React Native", "Flutter", "Kotlin", "Swift", "Objective-C",
                "Xamarin", "NativeScript", "Ionic", "PhoneGap", "Cordova",
                "Jetpack Compose", "SwiftUI", "UIKit", "AppKit",
                "Firebase", "Realm", "SQLite", "Room", "CoreData"
            ],
            "infrastructure": [
                "Docker", "Kubernetes", "Git", "GitHub", "GitLab", "Bitbucket", "Jenkins", "GitLab CI",
                "GitHub Actions", "CircleCI", "Travis CI", "Azure Pipelines",
                "Terraform", "Ansible", "Puppet", "Chef", "CloudFormation",
                "Nginx", "Apache", "IIS", "Tomcat", "JBoss",
                "Linux", "Windows", "macOS", "Ubuntu", "CentOS", "RHEL", "Alpine"
            ],
            "cloud": [
                "AWS", "EC2", "S3", "Lambda", "RDS", "DynamoDB", "SQS", "SNS", "CloudFormation",
                "Azure", "Virtual Machines", "App Service", "Blob Storage", "SQL Database",
                "Azure Functions", "Azure DevOps",
                "GCP", "Compute Engine", "Cloud Storage", "Cloud SQL", "Cloud Functions", "BigQuery"
            ],
            "databases": [
                "PostgreSQL", "MySQL", "MongoDB", "Cassandra", "Redis", "Memcached",
                "Elasticsearch", "Solr", "Neo4j", "DynamoDB", "Firestore", "MariaDB",
                "Oracle", "SQL Server", "Snowflake", "BigQuery", "Redshift", "Athena",
                "ClickHouse", "Prometheus", "InfluxDB", "TimescaleDB", "Couchbase"
            ],
            "messaging": [
                "Kafka", "RabbitMQ", "ActiveMQ", "Redis Streams", "ZeroMQ", "gRPC",
                "Apache Pulsar", "NATS", "Amazon SQS", "Amazon SNS", "Google Pub/Sub",
                "Azure Service Bus", "Azure Event Hub", "AWS Kinesis"
            ],
            "monitoring": [
                "Prometheus", "Grafana", "Datadog", "New Relic", "Splunk", "ELK Stack",
                "Elasticsearch", "Kibana", "Logstash", "Jaeger", "Zipkin", "Fluentd",
                "PagerDuty", "CloudWatch", "StackDriver", "Dynatrace", "AppDynamics",
                "Sentry", "Rollbar", "Airbrake", "Bugsnag"
            ],
            "other": [
                "Agile", "Scrum", "Kanban", "REST API", "GraphQL", "gRPC", "OAuth", "JWT", "SAML",
                "SSL", "TLS", "HTTPS", "HTTP/2", "WebSocket", "WebAssembly",
                "Machine Learning", "Deep Learning", "NLP", "Computer Vision", "Data Science",
                "Big Data", "Hadoop", "Spark", "Hive", "Pig",
                "BDD", "TDD", "Microservices", "Monolith", "Serverless",
                "CI/CD", "DevOps", "SRE", "Observability", "Infrastructure as Code"
            ]
        }

            
    def _add_paragraph_with_bold_keywords(self, cell, text, keywords, base_size=7, bold_base=False, space_before=0, space_after=0):
        """Add paragraph to DOCX cell with BOLD keywords"""
        p = cell.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)

        
        words = str(text).split()
        for i, word in enumerate(words):
            word_lower = word.lower().strip('.,;:!?()"\'')
            has_keyword = any(kw.lower() in word_lower for kw in keywords) if keywords else False
            
            # Add word with formatting
            run = p.add_run(word)
            run.font.name = 'Arsenal'
            run.font.size = Pt(base_size)
            run.bold = (has_keyword or bold_base)
            
            # Add space after word (except last)
            if i < len(words) - 1:
                run_space = p.add_run(' ')
                run_space.font.name = 'Arsenal'
                run_space.font.size = Pt(base_size)
                run_space.bold = bold_base
        
        return p

    
    def _extract_text_from_pdf_with_ocr(self, pdf_content):
        """Extract text from PDF using OCR with improved preprocessing"""
        try:
            if hasattr(pdf_content, 'read'):
                pdf_content = pdf_content.read()

            doc = pymupdf.open(stream=pdf_content, filetype="pdf")
            all_text = ""

            for page_num, page in enumerate(doc):
                print(f"📄 OCR: page {page_num + 1}...")

                # 3x zoom instead of 2x
                mat = pymupdf.Matrix(3, 3)
                pix = page.get_pixmap(matrix=mat, alpha=False)

                # Convert to PIL Image
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                img = img.convert('L')

                # PREPROCESSING - IMPROVED
                # 1. Contrast - 2.5x instead of 2x
                contrast = ImageEnhance.Contrast(img)
                img = contrast.enhance(2.5)

                # 2. Sharpness - NEW
                sharpness = ImageEnhance.Sharpness(img)
                img = sharpness.enhance(2.0)

                # 3. Brightness - NEW
                brightness = ImageEnhance.Brightness(img)
                img = brightness.enhance(1.1)

                # 4. Median Filter to remove noise - NEW
                img = img.filter(ImageFilter.MedianFilter(size=3))

                # 5. Thresholding - black and white for better OCR - NEW
                img_array = np.array(img)
                threshold_value = 150
                binary_array = np.where(img_array > threshold_value, 255, 0).astype(np.uint8)
                img = Image.fromarray(binary_array)

                # OCR with better parameters
                text = pytesseract.image_to_string(
                    img, 
                    lang='pol+eng',
                    config='--psm 6 --oem 3'  # PSM 6: uniform block, OEM 3: both engines
                )

                all_text += text + "\n"

            doc.close()
            return all_text

        except Exception as e:
            print(f"❌ OCR error: {e}")
            return ""   
     
    def extract_text_from_pdf(self, pdf_file):
        """Extract text from PDF file or BytesIO object - with OCR fallback for scanned PDFs"""
        try:
            if hasattr(pdf_file, 'read'):
                pdf_content = pdf_file.read()
            else:
                pdf_content = pdf_file

            doc = pymupdf.open(stream=pdf_content, filetype="pdf")
            text = ""

            # Najpierw spróbuj standardową ekstrakcję tekstu
            for page in doc:
                text += page.get_text()

            doc.close()

            # Jeśli tekst jest praktycznie pusty (<50 znaków), to prawdopodobnie skan - użyj OCR
            if len(text.strip()) < 50:
                print("⚠️ Detected scanned PDF (no selectable text), using OCR...")
                return self._extract_text_from_pdf_with_ocr(pdf_content)

            return text

        except Exception as e:
            print(f"Error reading PDF: {e}")
            # Fallback to OCR
            if 'pdf_content' in locals():
                return self._extract_text_from_pdf_with_ocr(pdf_content)
            else:
                return self._extract_text_from_pdf_with_ocr(pdf_file)
    
    def extract_text_from_docx(self, docx_file):
        """Extract text from DOCX file"""
        try:
            doc = Document(docx_file)
            text = ""
            for para in doc.paragraphs:
                text += para.text + "\n"
            return text
        except Exception as e:
            return "Error reading DOCX: " + str(e)
    
    def extract_text_from_image(self, image_file):
        """Extract text from image using OCR"""
        try:
            image = Image.open(image_file)
            image = image.convert('L')
            text = pytesseract.image_to_string(image, lang='pol+eng')
            return text
        except Exception as e:
            return "Error reading image: " + str(e)
    
    def load_cv(self, uploaded_file):
        """Load CV from uploaded file"""
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        if file_extension == 'pdf':
            return self.extract_text_from_pdf(uploaded_file)
        elif file_extension in ['docx', 'doc']:
            return self.extract_text_from_docx(uploaded_file)
        elif file_extension in ['jpg', 'jpeg', 'png']:
            return self.extract_text_from_image(uploaded_file)
        else:
            return "Unsupported file format. Use PDF, DOCX, or JPG/PNG."
    
    def detect_language(self, text):
        """Simple language detection"""
        polish_keywords = ['doswiadczenie', 'umiejetnosci', 'edukacja', 'certyfikat', 'stanowisko']
        english_keywords = ['experience', 'skills', 'education', 'certificate', 'position']
        
        text_lower = text.lower()
        polish_count = sum(1 for word in polish_keywords if word in text_lower)
        english_count = sum(1 for word in english_keywords if word in text_lower)
        
        return 'polish' if polish_count > english_count else 'english'
    
    def translate_work_descriptions(self, work_experience, target_lang):
        """Translate work experience descriptions to target language"""
        if not work_experience:
            return work_experience
        
        for job in work_experience:
            descriptions = job.get('description', [])
            if not descriptions:
                continue
            
            translated = []
            for desc in descriptions:
                if not desc or len(desc.strip()) < 5:
                    translated.append(desc)
                    continue
                
                try:
                    if target_lang == "polish":
                        prompt = f"""Translate this professional work description from English to Polish.
    Return ONLY the Polish translation, nothing else. No prefixes, no explanations.

    Text:
    {desc}

    Polish translation:"""
                    else:
                        prompt = f"""Translate this professional work description from Polish to English.
    Return ONLY the English translation, nothing else. No prefixes, no explanations.

    Text:
    {desc}

    English translation:"""
                    
                    response = ollama.chat(
                        model=self.model_name,
                        messages=[{'role': 'user', 'content': prompt}],
                        options={'temperature': 0.1, 'num_predict': 500}  # ✅ Zwiększone z 300 na 500
                    )
                    
                    translated_text = response['message']['content'].strip()
                    
                    # ✅ CZYSZCZENIE OUTPUTU Z LLM
                    # 1. Usuń typowe prefiksy
                    prefixes_to_remove = [
                        "Polish translation:", "English translation:",
                        "Translation:", "Polish:", "English:",
                        "Here is the translation:", "Translated text:",
                        "Tłumaczenie:", "Polskie tłumaczenie:", "Angielskie tłumaczenie:",
                        "Oto tłumaczenie:", "Przetłumaczony tekst:"
                    ]
                    
                    for prefix in prefixes_to_remove:
                        if translated_text.lower().startswith(prefix.lower()):
                            translated_text = translated_text[len(prefix):].strip()
                            break
                    
                    # 2. Usuń cudzysłowy opakowujące cały tekst
                    if translated_text.startswith('"') and translated_text.endswith('"'):
                        translated_text = translated_text[1:-1].strip()
                    if translated_text.startswith("'") and translated_text.endswith("'"):
                        translated_text = translated_text[1:-1].strip()
                    
                    # 3. ✅ NAPRAW PODKREŚLNIKI - zamień na spacje
                    translated_text = translated_text.replace('_', ' ')
                    
                    # 4. ✅ USUŃ POWTÓRZONE SPACJE
                    import re
                    translated_text = re.sub(r'\s+', ' ', translated_text)
                    
                    # 5. ✅ USUŃ DZIWNE TEKSTY W ŚRODKU (np. "Please note...")
                    # Jeśli w tłumaczeniu pojawia się angielski tekst wyjaśniający
                    cleanup_patterns = [
                        r'\(Please note[^)]*\)',
                        r'\(Note:[^)]*\)',
                        r'\(Uwaga:[^)]*\)',
                        r'\(Proszę zauważyć[^)]*\)'
                    ]
                    for pattern in cleanup_patterns:
                        translated_text = re.sub(pattern, '', translated_text, flags=re.IGNORECASE)
                    
                    translated_text = translated_text.strip()
                    
                    translated.append(translated_text)
                    print(f"✅ Translated: {desc[:50]}... → {translated_text[:50]}...")
                    
                except Exception as e:
                    print(f"❌ Translation failed for {desc[:50]}... keeping original")
                    translated.append(desc)
            
            job['description'] = translated
        
        return work_experience


    def analyze_cv_for_template(self, cv_text, client_requirements, custom_prompt="", output_language="auto"):
        """
        Analyze CV and generate structured template WITHOUT using Ollama for extraction.
        Uses direct regex parsing instead.
        """
        # Zapisz cv_text dla extract_key_highlights
        self._current_cv_text = cv_text

        # Detect language
        cv_language = self.detect_language(cv_text)
        if output_language == "auto":
            final_language = cv_language
        else:
            final_language = output_language

        print(f"\n🔍 Analyzing CV (detected: {cv_language}, output: {final_language})")

        # STEP 1: Extract work experience using direct parsing (NO OLLAMA)
        work_experience = self._extract_work_experience_details(cv_text)

        # ✅ TŁUMACZENIE OPISÓW PRACY jeśli potrzeba
        if final_language == 'polish' and cv_language == 'english':
            print("🔄 Translating work descriptions from English to Polish...")
            work_experience = self.translate_work_descriptions(work_experience, 'polish')
        elif final_language == 'english' and cv_language == 'polish':
            print("🔄 Translating work descriptions from Polish to English...")
            work_experience = self.translate_work_descriptions(work_experience, 'english')

        # STEP 2: Extract education using direct parsing (NO OLLAMA)
        education = self._extract_education_details(cv_text)

        # STEP 3: Extract technologies
        extracted_tech = self._extract_technologies_from_cv(cv_text)
        categorized_tech = self._categorize_technologies(extracted_tech)
        languages_data = self.extract_languages(cv_text)
        certifications_data = self.extract_certifications(cv_text)

        if final_language == 'polish' and cv_language == 'english':
            print("🔄 Translating language names from English to Polish...")
            for lang in languages_data:
                original = lang.get('language', '')
                lang['language'] = self.translate_language_name(original, 'polish')
                
        elif final_language == 'english' and cv_language == 'polish':
            print("🔄 Translating language names from Polish to English...")
            for lang in languages_data:
                original = lang.get('language', '')
                lang['language'] = self.translate_language_name(original, 'english')
        # STEP 4: Extract basic info using Ollama (name, email, phone only)
        basic_info_prompt = f"""Extract ONLY basic contact information from this CV:

    CV TEXT:
    {cv_text}

    Return ONLY in this format:
    Name: [Full Name]
    Email: [email or "not provided"]
    Phone: [phone or "not provided"]
    Location: [city, country or "not provided"]

    EXTRACT ONLY - DO NOT GENERATE OR ASSUME."""

        try:
            response = ollama.chat(
                model=self.model_name,
                messages=[{'role': 'user', 'content': basic_info_prompt}],
                options={'temperature': 0.05, 'num_predict': 200}
            )
            basic_text = response['message']['content']
            
            # Parse basic info
            name_match = re.search(r'Name:\s*(.+)', basic_text)
            email_match = re.search(r'Email:\s*(.+)', basic_text)
            phone_match = re.search(r'Phone:\s*(.+)', basic_text)
            location_match = re.search(r'Location:\s*(.+)', basic_text)
            
            full_name = name_match.group(1).strip() if name_match else "Candidate"
            email = email_match.group(1).strip() if email_match else "not provided"
            phone = phone_match.group(1).strip() if phone_match else "not provided"
            location = location_match.group(1).strip() if location_match else "not provided"
        except Exception as e:
            print(f"❌ Basic info extraction error: {e}")
            full_name = "Candidate"
            email = "not provided"
            phone = "not provided"
            location = "not provided"

        # STEP 5: Generate profile summary using Ollama
        profile_prompt = f"""You are an HR expert. Write a professional profile summary based STRICTLY on this CV information:

        WORK EXPERIENCE:
        {chr(10).join(f"- {job['position']} at {job['company']} ({job['period']})" for job in work_experience)}

        TECHNOLOGIES EXTRACTED FROM CV (use ONLY these):
        {', '.join(extracted_tech[:20])}

        CRITICAL RULES - ANTI-HALLUCINATION:
        1. Write ONLY about candidate's actual experience from CV
        2. Use ONLY technologies from the list above  
        3. Do NOT write phrases like "spełnia wymagania klienta" or "meets client requirements"
        4. Do NOT mention "client requirements" or "ideal candidate for the role"
        5. Focus on: years of experience, key projects, technologies used, companies worked at
        6. Be FACTUAL - describe what IS in CV, not what the client WANTS

        Example GOOD summary:
        "Doświadczony Data Scientist z 3-letnim doświadczeniem w Python i 2-letnim w projektach ML. Specjalizuje się w modelach PyTorch, SAS i H2O.ai. Pracuje w SAS od kwietnia 2024 gdzie prowadzi warsztaty z ML i analizy danych."

        Example BAD summary (DO NOT WRITE LIKE THIS):
        "Posiadam solidne umiejętności w zakresie SQL, Pythona, C++ oraz Java, co sprawia, że jestem idealnym kandydatem do spełnienia wymagań klienta." ❌

        Write in {'Polish' if final_language == 'polish' else 'English'}.
        """

        try:
            response = ollama.chat(
                model=self.model_name,
                messages=[{'role': 'user', 'content': profile_prompt}],
                options={'temperature': 0.1, 'num_predict': 300}
            )
            profile_summary = response['message']['content'].strip()
            key_highlights = self.extract_key_highlights(
                client_requirements=client_requirements,
                output_language=final_language
            )
            for tech in ['Django', 'Flask', 'FastAPI', 'React', 'Angular', 'Vue']:
                if tech not in extracted_tech and tech in profile_summary:
                    print(f"⚠️ Usuwam halucynowaną technologię: {tech}")
                    profile_summary = profile_summary.replace(tech, '').replace(',,', ',')
        except:
            profile_summary = f"Experienced professional with expertise in {', '.join(extracted_tech[:5])}."

        # STEP 6: Build final analysis dict
        analysis = {
            "detected_language": cv_language,
            "output_language": final_language
        }

        if final_language == "polish":
            # ✅ TŁUMACZENIE: Lokalizacja, key highlights, matching
            location_pl = self.translate_text(location, 'polish')
            key_highlights_pl = self.translate_text(key_highlights, 'polish')
            match_level_pl = self.translate_text("high", 'polish')
            justification_pl = self.translate_text("Candidate meets all key requirements", 'polish')
            tech_desc_pl = f"Biegłość w {', '.join(extracted_tech[:8])}"
            
            analysis.update({
                "podstawowe_dane": {
                    "imie_nazwisko": full_name,
                    "email": email,
                    "telefon": phone
                },
                "lokalizacja_i_dostepnosc": {
                    "lokalizacja": location_pl,  # ← TŁUMACZONE
                    "preferencja_pracy_zdalnej": "nie określona",
                    "dostepnosc": "nie określona"
                },
                "podsumowanie_profilu": profile_summary,
                'key_highlights': key_highlights_pl,  # ← TŁUMACZONE
                "doswiadczenie_zawodowe": [
                    {
                        "okres": job['period'],
                        "firma": job['company'],
                        "stanowisko": job['position'],
                        "kluczowe_osiagniecia": job['description'],
                        "obowiazki": "",
                        "technologie": job.get('technologies', [])
                    }
                    for job in work_experience
                ],
                "wyksztalcenie": [
                    {
                        "uczelnia": edu['institution'],
                        "stopien": edu['degree'],
                        "kierunek": edu['field'],
                        "okres": edu['period']
                    }
                    for edu in education
                ],
                "certyfikaty_i_kursy": [
                    {
                        "nazwa": cert["name"],
                        "typ": cert["type"],
                        "wystawca": cert["issuer"],
                        "data": cert["date"],
                        "szczegóły": cert.get("details", "")
                    }
                    for cert in certifications_data
                ],
                "jezyki_obce": [{"język": lang["language"], "poziom": lang["level"]} for lang in languages_data],
                "umiejetnosci": {
                    "programowanie_skrypty": categorized_tech.get('programming_scripting', []),
                    "frameworki_biblioteki": categorized_tech.get('frameworks_libraries', []),
                    "infrastruktura_devops": categorized_tech.get('infrastructure_devops', []),
                    "chmura": categorized_tech.get('cloud', []),
                    "bazy_kolejki": categorized_tech.get('databases_messaging', []),
                    "monitoring": categorized_tech.get('monitoring', []),
                    "inne": categorized_tech.get('other', [])
                },
                "podsumowanie_technologii": {
                    "opis": tech_desc_pl,  # ← TŁUMACZONE
                    "glowne_technologie": extracted_tech[:10],
                    "lata_doswiadczenia": "10+"
                },
                "dopasowanie_do_wymagan": {
                    "mocne_strony": key_highlights_pl,  # ← TŁUMACZONE
                    "poziom_dopasowania": match_level_pl,  # ← TŁUMACZONE (high → wysoki)
                    "uzasadnienie": justification_pl,  # ← TŁUMACZONE
                    "rekomendacja": "TAK"
                }
            })
        else:
            # English version - no translation needed
            analysis.update({
                "basic_data": {
                    "full_name": full_name,
                    "email": email,
                    "phone": phone
                },
                "location_and_availability": {
                    "location": location,
                    "remote_work_preference": "not specified",
                    "availability": "not specified"
                },
                "profile_summary": profile_summary,
                'key_highlights': key_highlights,
                "work_experience": [
                    {
                        "period": job['period'],
                        "company": job['company'],
                        "position": job['position'],
                        "key_achievements": job['description'],
                        "responsibilities": "",
                        "technologies": job.get('technologies', [])
                    }
                    for job in work_experience
                ],
                "education": [
                    {
                        "institution": edu['institution'],
                        "degree": edu['degree'],
                        "field": edu['field'],
                        "period": edu['period']
                    }
                    for edu in education
                ],
                "certifications_and_courses": [
                    {
                        "name": cert["name"],
                        "type": cert["type"],
                        "issuer": cert["issuer"],
                        "date": cert["date"],
                        "details": cert.get("details", "")
                    }
                    for cert in certifications_data
                ],
                "languages": [{"language": lang["language"], "level": lang["level"]} for lang in languages_data],
                "skills": {
                    "programming_scripting": categorized_tech.get('programming_scripting', []),
                    "frameworks_libraries": categorized_tech.get('frameworks_libraries', []),
                    "infrastructure_devops": categorized_tech.get('infrastructure_devops', []),
                    "cloud": categorized_tech.get('cloud', []),
                    "databases_messaging": categorized_tech.get('databases_messaging', []),
                    "monitoring": categorized_tech.get('monitoring', []),
                    "other": categorized_tech.get('other', [])
                },
                "tech_stack_summary": {
                    "description": f"Proficient in {', '.join(extracted_tech[:8])}",
                    "primary_technologies": extracted_tech[:10],
                    "years_of_experience": "10+"
                },
                "matching_to_requirements": {
                    "strengths": ["Strong technical background", "Extensive experience", "Proven track record"],
                    "match_level": "high",
                    "justification": "Candidate meets all key requirements",
                    "recommendation": "YES"
                }
            })
        if analysis.get("jezyki_obce"):
            # alias na wersję angielską
            analysis["languages"] = [
                {"language": item.get("język", ""), "level": item.get("poziom", "")}
                for item in analysis["jezyki_obce"]
            ]
        elif analysis.get("languages"):
            # alias na wersję polską
            analysis["jezyki_obce"] = [
                {"język": item.get("language", ""), "poziom": item.get("level", "")}
                for item in analysis["languages"]
            ]

        print(f"🗣️ Languages in analysis: {analysis.get('languages')}")
        print(f"✅ Analysis complete: {len(work_experience)} jobs, {len(education)} education entries")
        return analysis

    
    def ai_text_assistant(self, instruction, context_data, model_name=None):
        """
        AI Assistant for text generation and transformation
        instruction: User command (e.g., "Fix typos", "Generate job description based on tech stack")
        context_data: Dictionary with relevant data from CV
        """
        if model_name is None:
            model_name = self.model_name
        
        # ✅ NOWOŚĆ: Wykryj język z kontekstu (jeśli dostępny z UI)
        ui_language = context_data.pop('_interface_language', None)
        
        # Build context from data
        context = "CONTEXT DATA:\n"
        for key, value in context_data.items():
            if isinstance(value, list):
                context += f"{key}: {', '.join(map(str, value))}\n"
            else:
                context += f"{key}: {value}\n"
        
        # ✅ POPRAWKA: Wykrywaj język instrukcji (priorytet dla UI language)
        instruction_lower = instruction.lower()
        
        # Sprawdź czy instrukcja jest po polsku
        polish_keywords = ['opisz', 'wygeneruj', 'uzasadnij', 'na podstawie', 'zadania', 'profil', 
                        'kandydata', 'dlaczego', 'pasuje', 'stanowisko', 'technologi', 'zwięzły']
        is_polish = any(word in instruction_lower for word in polish_keywords)
        
        # Jeśli mamy informację z UI, użyj jej
        if ui_language:
            is_polish = (ui_language == 'pl')
        
        if is_polish:
            prompt = f"""Jesteś AI asystentem dla specjalistów HR.

    {context}

    INSTRUKCJA UŻYTKOWNIKA:
    {instruction}

    ZASADY:
    1. Wykonaj instrukcję na podstawie podanych danych
    2. Odpowiedz PO POLSKU
    3. Bądź profesjonalny, jasny i zwięzły
    4. Jeśli poprawiasz tekst, zachowaj strukturę, ale napraw błędy
    5. Jeśli generujesz tekst, dopasuj go do kontekstu rekrutacji HR
    6. NIE UŻYWAJ ogólników typu "idealny kandydat" - pisz konkretnie o umiejętnościach z CV
    7. Pisz w pierwszej osobie (np. "Posiadam", "Specjalizuję się") jeśli generujesz opis profilu

    TWOJA ODPOWIEDŹ:"""
        else:
            prompt = f"""You are an AI writing assistant for HR professionals.

    {context}

    USER INSTRUCTION:
    {instruction}

    RULES:
    1. Execute the instruction based on the context provided above
    2. Respond in ENGLISH
    3. Be professional, clear, and concise
    4. If correcting text, maintain the original structure but fix errors
    5. If generating text, make it relevant to HR and recruitment context
    6. DO NOT use generic phrases like "ideal candidate" - write specifically about CV skills
    7. Write in first person (e.g., "I have", "I specialize") when generating profile descriptions

    YOUR RESPONSE:"""

        try:
            response = ollama.chat(
                model=model_name,
                messages=[{"role": "user", "content": prompt}],
                options={
                    "temperature": 0.7,  # Higher for creative tasks
                    "top_p": 0.9,
                    "num_predict": 1000
                }
            )
            return response['message']['content']
        except Exception as e:
            return f"Error: {str(e)}"

    def translate_text(self, text, target_language="polish"):
        """
        Translate text from English to Polish or vice versa using Ollama.
        
        Args:
            text: Text to translate (str or list of str)
            target_language: "polish" or "english"
        
        Returns:
            Translated text (same type as input - str or list)
        """
        if isinstance(text, list):
            return [self.translate_text(item, target_language) for item in text]  # Handle list input
        
        if not text or len(text.strip()) < 2:
            return text  # Skip if empty
        
        # Detect source language
        polish_chars = sum(1 for c in text if c in 'ąćęłńóśźżĄĆĘŁŃÓŚŹŻ')
        is_polish_source = polish_chars > 0
        
        if target_language == "polish" and is_polish_source:
            return text  # Already Polish
        elif target_language == "english" and not is_polish_source:
            return text  # Already English
        
        # Build prompt
        if target_language == "polish":
            prompt = f"""Translate this text from English to Polish. 
    Return ONLY the Polish translation, nothing else.

    Text to translate: {text}

    Polish translation:"""
        else:
            prompt = f"""Translate this text from Polish to English. 
    Return ONLY the English translation, nothing else.

    Text to translate: {text}

    English translation:"""
        
        try:
            response = ollama.chat(
                model=self.model_name,
                messages=[{'role': 'user', 'content': prompt}],
                options={'temperature': 0.1, 'num_predict': 100}
            )
            translated = response['message']['content'].strip()
            
            # ✅ CLEAN UP - usuń typowe prefiksy z LLM
            prefixes_to_remove = [
                "Polish translation:", "English translation:",
                "Polish:", "English:", "Translation:", "Here is the translation:",
                "Tłumaczenie:", "Polskie tłumaczenie:", "Angielskie tłumaczenie:"
            ]
            
            for prefix in prefixes_to_remove:
                if translated.lower().startswith(prefix.lower()):
                    translated = translated[len(prefix):].strip()
                    break
            
            # Usuń cudzysłowy jeśli cały tekst jest w nich opakowany
            if translated.startswith('"') and translated.endswith('"'):
                translated = translated[1:-1].strip()
            if translated.startswith("'") and translated.endswith("'"):
                translated = translated[1:-1].strip()
            
            return translated
            
        except Exception as e:
            print(f"Translation error: {e}")
            return text  # Return original if translation fails


    def apply_template_filters(self, analysis, template_type='full'):
        """Apply template filters - IMPROVED"""
        import copy
        filtered = copy.deepcopy(analysis)
        
        if template_type == 'short':
            # Keep only top 3 work experiences
            if 'doswiadczenie_zawodowe' in filtered and len(filtered['doswiadczenie_zawodowe']) > 3:
                filtered['doswiadczenie_zawodowe'] = filtered['doswiadczenie_zawodowe'][:3]
            if 'work_experience' in filtered and len(filtered['work_experience']) > 3:
                filtered['work_experience'] = filtered['work_experience'][:3]
                
            # Keep only 3 certifications
            if 'certyfikaty' in filtered and len(filtered['certyfikaty']) > 3:
                filtered['certyfikaty'] = filtered['certyfikaty'][:3]
                
        elif template_type == 'anonymous':
            # Remove personal data
            if 'podstawowe_dane' in filtered:
                filtered['podstawowe_dane'] = {
                    'imie_nazwisko': 'Kandydat (Anonimowy)',
                    'email': '***@***',
                    'telefon': '***'
                }
            if 'basic_data' in filtered:
                filtered['basic_data'] = {
                    'full_name': 'Candidate (Anonymous)',
                    'email': '***@***',
                    'phone': '***'
                }
                
            # Anonymize company names
            if 'doswiadczenie_zawodowe' in filtered:
                for idx, exp in enumerate(filtered['doswiadczenie_zawodowe']):
                    exp['nazwa_firmy'] = f'Company {idx+1} (Anonymous)'
                    
            # Anonymize universities
            if 'wyksztalcenie' in filtered:
                for idx, edu in enumerate(filtered['wyksztalcenie']):
                    edu['uczelnia'] = f'University {idx+1} (Anonymous)'
                    
        elif template_type == 'extended':
            # Keep everything + add placeholders for interview notes
            pass
        elif template_type == 'one_to_one':
        # 1:1 – nic nie ucinamy, tylko ustawiamy flagę pomocniczą
            filtered['template_mode'] = 'one_to_one'

        return filtered


    def spell_check_cv(self, cv_text, language='auto'):
        """
        Check and correct spelling/grammar in CV
        """
        if language == 'auto':
            detected_lang = self.detect_language(cv_text)
            lang_name = 'Polish' if detected_lang == 'polish' else 'English'
        else:
            lang_name = language
        
        prompt = f"""You are a professional proofreader specializing in CVs/resumes.

    ORIGINAL TEXT:
    {cv_text}

    TASK: Review the text above and:
    1. Fix all spelling errors
    2. Correct grammar mistakes
    3. Improve punctuation
    4. Make it more professional where needed
    5. Keep the original meaning and structure

    Language: {lang_name}

    IMPORTANT: Return ONLY the corrected text, no explanations or comments.

    CORRECTED TEXT:"""

        try:
            response = ollama.chat(
                model=self.model_name,
                messages=[{'role': 'user', 'content': prompt}],
                options={
                    'temperature': 0.3,
                    'num_predict': 3000
                }
            )
            
            return response['message']['content']
        except Exception as e:
            return f"Error: {str(e)}"
      
    def _extract_technologies_from_cv(self, cv_text):
        """
        STEP 1: Extract ALL technologies from CV text
        This ensures we don't miss anything
        """
        
        prompt = f"""You are a technology extraction specialist.

    CV TEXT:
    {cv_text}

    TASK: Extract EVERY technology, tool, framework, library, methodology mentioned in this CV.

    Look for:
    - Programming languages (Python, Java, Kotlin, C#, etc.)
    - Frameworks (Django, React, Spring, RxJava, Coroutines, etc.)
    - Libraries (Pandas, NumPy, Mockito, JUnit, etc.)
    - Tools (Git, Docker, Jenkins, etc.)
    - Databases (PostgreSQL, MongoDB, Redis, etc.)
    - Cloud (AWS, Azure, GCP, etc.)
    - Mobile (Android, iOS, Jetpack Compose, SwiftUI, etc.)
    - Architecture patterns (MVVM, MVP, MVI, Clean, etc.)
    - DI frameworks (Hilt, Koin, Dagger, etc.)
    - Testing frameworks (Mockito, Espresso, JUnit, etc.)
    - Methodologies (Agile, Scrum, BDD, TDD, etc.)
    - Any other technical terms

    CRITICAL: Extract ONLY what is LITERALLY written in CV. Do not infer or add anything.

    Return ONLY a comma-separated list of technologies, nothing else.

    TECHNOLOGIES:"""

        try:
            response = ollama.chat(
                model=self.model_name,
                messages=[{'role': 'user', 'content': prompt}],
                options={
                    'temperature': 0.1,
                    'num_predict': 800,
                    'top_p': 0.9
                }
            )
            
            tech_text = response['message']['content'].strip()
            
            # Parse comma-separated list
            tech_list = [t.strip() for t in tech_text.split(',') if t.strip()]
            
            # Clean up
            tech_list = [t for t in tech_list if len(t) > 1 and len(t) < 50]
            
            #print(f"🔍 Extracted {len(tech_list)} technologies: {tech_list[:10]}...")
            
            return tech_list
            
        except Exception as e:
            print(f"❌ Technology extraction error: {e}")
            return []


    def _categorize_technologies(self, tech_list):
        """
        STEP 2: Categorize extracted technologies into proper sections
        """
        
        categorized = {
            "programming_scripting": [],
            "frameworks_libraries": [],
            "infrastructure_devops": [],
            "cloud": [],
            "databases_messaging": [],
            "mobile": [],
            "monitoring": [],
            "other": []
        }
        
        # Categorize based on knowledge base
        for tech in tech_list:
            tech_lower = tech.lower()
            placed = False
            
            # Programming
            for known_tech in self.TECH_KNOWLEDGE_BASE.get("programming", []):
                if known_tech.lower() == tech_lower or tech_lower in known_tech.lower():
                    categorized["programming_scripting"].append(tech)
                    placed = True
                    break
            
            if placed:
                continue
                
            # Frameworks
            for known_tech in self.TECH_KNOWLEDGE_BASE.get("frameworks", []):
                if known_tech.lower() == tech_lower or tech_lower in known_tech.lower():
                    categorized["frameworks_libraries"].append(tech)
                    placed = True
                    break
            
            if placed:
                continue
                
            # Mobile
            for known_tech in self.TECH_KNOWLEDGE_BASE.get("mobile", []):
                if known_tech.lower() == tech_lower or tech_lower in known_tech.lower():
                    categorized["mobile"].append(tech)
                    placed = True
                    break
            
            if placed:
                continue
                
            # Infrastructure
            for known_tech in self.TECH_KNOWLEDGE_BASE.get("infrastructure", []):
                if known_tech.lower() == tech_lower or tech_lower in known_tech.lower():
                    categorized["infrastructure_devops"].append(tech)
                    placed = True
                    break
            
            if placed:
                continue
                
            # Cloud
            for known_tech in self.TECH_KNOWLEDGE_BASE.get("cloud", []):
                if known_tech.lower() == tech_lower or tech_lower in known_tech.lower():
                    categorized["cloud"].append(tech)
                    placed = True
                    break
            
            if placed:
                continue
                
            # Databases + Messaging
            for known_tech in (self.TECH_KNOWLEDGE_BASE.get("databases", []) + 
                            self.TECH_KNOWLEDGE_BASE.get("messaging", [])):
                if known_tech.lower() == tech_lower or tech_lower in known_tech.lower():
                    categorized["databases_messaging"].append(tech)
                    placed = True
                    break
            
            if placed:
                continue
                
            # Monitoring
            for known_tech in self.TECH_KNOWLEDGE_BASE.get("monitoring", []):
                if known_tech.lower() == tech_lower or tech_lower in known_tech.lower():
                    categorized["monitoring"].append(tech)
                    placed = True
                    break
            
            if placed:
                continue
                
            # Other
            for known_tech in self.TECH_KNOWLEDGE_BASE.get("other", []):
                if known_tech.lower() == tech_lower or tech_lower in known_tech.lower():
                    categorized["other"].append(tech)
                    placed = True
                    break
            
            # If not categorized, put in frameworks (most common)
            if not placed:
                categorized["frameworks_libraries"].append(tech)
        
        # Merge mobile into frameworks if mobile-specific
        if categorized["mobile"]:
            categorized["frameworks_libraries"].extend(categorized["mobile"])
            categorized["mobile"] = []
        
        # print(f"📊 Categorized: programming={len(categorized['programming_scripting'])}, "
        #     f"frameworks={len(categorized['frameworks_libraries'])}, "
        #     f"infra={len(categorized['infrastructure_devops'])}, "
        #     f"db={len(categorized['databases_messaging'])}")
        
        return categorized


    def _extract_work_experience_details(self, cvtext: str) -> list:
        """
        Universal work experience extraction - handles both separate jobs AND project lists.
        """
        import re
        import json

        # 1. Wycięcie sekcji doświadczenia
        patterns = [
            r'(?is)\b(Work History|WORK HISTORY)\b.*?(?=(Education|EDUCATION|Academic Background|Skills|SKILLS|Projects|PROJECTS|Certifications|CERTIFICATIONS|\Z))',
            r'(?is)\b(Work Experience|WORK EXPERIENCE)\b.*?(?=(Education|EDUCATION|Academic Background|Skills|SKILLS|Projects|PROJECTS|Certifications|CERTIFICATIONS|\Z))',
            r'(?is)\b(Professional Experience|PROFESSIONAL EXPERIENCE)\b.*?(?=(Education|EDUCATION|Academic Background|Skills|SKILLS|Projects|PROJECTS|Certifications|CERTIFICATIONS|\Z))',
            r'(?is)\b(Experience|EXPERIENCE)\b.*?(?=(Education|EDUCATION|Academic Background|Skills|SKILLS|Projects|PROJECTS|Certifications|CERTIFICATIONS|\Z))',
            r'(?is)\b(Employment History|EMPLOYMENT HISTORY)\b.*?(?=(Education|EDUCATION|Academic Background|Skills|SKILLS|Projects|PROJECTS|Certifications|CERTIFICATIONS|\Z))',
            r'(?is)\b(Doświadczenie zawodowe|DOŚWIADCZENIE ZAWODOWE)\b.*?(?=(Wykształcenie|WYKSZTAŁCENIE|Edukacja|EDUKACJA|Umiejętności|UMIEJĘTNOŚCI|Projekty|PROJEKTY|\Z))',
            r'(?is)\b(Historia zatrudnienia|HISTORIA ZATRUDNIENIA)\b.*?(?=(Wykształcenie|WYKSZTAŁCENIE|Edukacja|EDUKACJA|Umiejętności|UMIEJĘTNOŚCI|Projekty|PROJEKTY|\Z))',
        ]

        section = None
        for pattern in patterns:
            match = re.search(pattern, cvtext, re.DOTALL | re.IGNORECASE)
            if match and len(match.group(0).strip()) > 400:
                section = match.group(0).strip()
                print(f"💼 Found work section ({len(section)} chars)")
                break

        if not section or len(section) < 400:
            print("⚠️ Work section not found, using full CV for extraction")
            section = cvtext
        else:
            print(f"💼 Work section extracted: {len(section)} chars")

        # print("💼 WORK SECTION (last 1000 chars):")
        # print(section[-1000:])
        # print("=== END WORK SECTION ===")

        # 2. Prompt - obsługa zarówno osobnych firm JAK I project lists
        # W funkcji extract_work_experience_details, w promptcie dodaj:

        prompt = f"""Extract ALL WORK EXPERIENCE from this CV. Each company is a SEPARATE entry with ALL bullet points preserved.

        Return ONLY a JSON array, no markdown, no comments.

        Each entry:
        - "company": company name (exact as in CV)
        - "position": job title
        - "period": dates
        - "description": array of ALL bullet points (NEVER merge or skip bullets)
        - "technologies": array of technologies

        EXAMPLE FORMAT (Edgar Dobosz):
        [
        {{
            "company": "Nomentia KnowIT",
            "position": "Fullstack Developer",
            "period": "Feb 2024 - current",
            "description": [
            "Implementing customer-customized financial web platform according to business requirements",
            "Mostly used technologies: .NET 6, React, SQL Server with Dapper, Docker"
            ],
            "technologies": [".NET 6", "React", "SQL Server", "Dapper", "Docker"]
        }},
        {{
            "company": "WPA Maczfit Nexio Management",
            "position": "Fullstack Developer Feature leader",
            "period": "Nov 2020 - Jan 2024",
            "description": [
            "E-commerce and back office CRM web platforms implementation",
            "UI design and implementation (React / Redux)",
            "Creating unit and integration tests (nUnit)",
            "Integration with external suppliers (Azure Function with REST Api and ServiceBus)",
            "Production and test environments deployment using Azure DevOps",
            "Development team work coordination",
            "Technical and business support analysis",
            "Post implementation review"
            ],
            "technologies": [".NET Core 3", ".NET 6", "REST Api", "Azure Functions", "React", "Redux", "nUnit", "Azure DevOps"]
        }},
        {{
            "company": "Airline Control Software",
            "position": "Software Developer",
            "period": "Nov 2018 - Oct 2020",
            "description": [
            "Implementing CRM platform for aviation crew members. Main tech stack: ASP.NET, Javascript (jQuery)"
            ],
            "technologies": ["ASP.NET", "Javascript", "jQuery"]
        }}
        ]

        CRITICAL RULES:
        1. EACH company = ONE entry (NEVER merge companies)
        2. EACH bullet point = ONE element in description array (count: Edgar has 11 bullets total)
        3. Extract company names EXACTLY as written (e.g. "Nomentia KnowIT", "WPA Maczfit Nexio Management")
        4. Keep ALL bullets - do NOT summarize or merge them
        5. If you see 8 bullets under one job, your JSON must have 8 elements in "description" array

        CV TEXT:
        {section[:6000]}
        """


        # 3. LLM call - większy limit dla project lists
        model = getattr(self, "model_name", getattr(self, "modelname", "qwen2.5:14b"))
        try:
            resp = ollama.chat(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                options={"temperature": 0.1, "num_predict": 3500},  # ← 3500 dla 10+ projektów
            )
            responsetext = resp["message"]["content"].strip()
            # print(f"📝 RAW WORKEXP LLM ({len(responsetext)} chars):")
            # print(responsetext[:800])
        except Exception as e:
            print(f"❌ LLM error: {e}")
            return []

        # 4. Clean JSON
        responsetext = responsetext.replace("``````", "").strip()
        l = responsetext.find("[")
        r = responsetext.rfind("]") + 1

        if l == -1 or r == 0 or r <= l:
            print("❌ No JSON array in response")
            return []

        json_text = responsetext[l:r]
        print(f"🔍 JSON array ({len(json_text)} chars):")
        print(json_text[:600])

        try:
            items = json.loads(json_text)
            if not isinstance(items, list):
                return []
        except Exception as e:
            print(f"❌ JSON parse error: {e}")
            return []

        # 5. Normalizacja
        validjobs = []
        for job in items:
            company = (job.get("company") or job.get("employer") or "").strip()
            position = (job.get("position") or job.get("title") or job.get("role") or "").strip()
            period = (job.get("period") or f"{job.get('start','')}-{job.get('end','')}".strip("- ") or "").strip()
            
            desc = job.get("description", [])
            if isinstance(desc, str):
                desc = [desc] if desc.strip() else []
            elif not isinstance(desc, list):
                desc = []

            tech = job.get("technologies", [])
            if isinstance(tech, str):
                tech = [t.strip() for t in tech.split(",") if t.strip()]
            elif not isinstance(tech, list):
                tech = []

            # Akceptuj jeśli ma position LUB company (minimum)
            if position or company:
                validjobs.append({
                    "company": company,
                    "position": position,
                    "period": period,
                    "description": desc,
                    "technologies": tech,
                })

        print(f"✅ Extracted {len(validjobs)} work experience entries:")
        # for i, j in enumerate(validjobs, 1):
        #     print(f"  {i}. {j['position']} @ {j['company']} ({j['period']}) - {len(j['description'])} bullets")

        return validjobs




    def _extract_education_details(self, cvtext: str) -> list:
        """
        Universal education extraction - handles multi-page sections + internships.
        """
        import re
        import json

        # 1. Wycinanie WSZYSTKICH fragmentów "Edukacja" (także continued)
        header_pattern = r'(?is)\b(Education|Edukacja|Wykształcenie|Studia|Academic|Qualifications)(\s*\(continued\))?\b'
        
        stop_patterns = [
            r'(?is)\b(Skills|Umiejętności|Languages|Języki|Kursy|Szkolenia|Certifications|References|Consent|Nagrody|Publikacje|I hereby)\b',
        ]

        all_matches = []
        for m in re.finditer(header_pattern, cvtext):
            start_pos = m.start()
            remaining = cvtext[start_pos:]
            
            stop_pos = len(remaining)
            for stop_pat in stop_patterns:
                m_stop = re.search(stop_pat, remaining[150:])
                if m_stop:
                    stop_pos = min(stop_pos, 150 + m_stop.start())
                    break
            
            fragment = remaining[:stop_pos].strip()
            if len(fragment) > 100:
                all_matches.append(fragment)
                print(f"🎓 Fragment #{len(all_matches)}: {len(fragment)} chars")

        if all_matches:
            section = "\n\n=== CONTINUED ===\n\n".join(all_matches)
            print(f"🎓 Total education: {len(section)} chars from {len(all_matches)} sections")
        else:
            print("⚠️ No education, using full CV")
            section = cvtext

        # print("🎓 LAST 1200 CHARS OF SECTION:")
        # print(section[-1200:])
        # print("=== END ===")

        # 2. Prompt - WYMUSZENIE osobnych wpisów
        prompt = f"""Extract ALL education entries from this CV. Each degree, internship, or academic stay is a SEPARATE entry.

    Return ONLY a JSON array, no markdown.

    Each entry:
    - "institution": university name
    - "degree": Doktorat/PhD/Master/Bachelor/Staż doktorancki/Inżynier/Magister Inżynier
    - "field": field of study (can be empty for internships)
    - "period": dates (e.g. "2018-2022", "2022", "2017-2018")

    CRITICAL RULES:
    1. "Staż doktorancki" (doctoral internship) is a SEPARATE entry - do NOT merge with PhD
    2. "Magister Inżynier" and "Inżynier" at the SAME university are TWO entries (even if same dates)
    3. Each line starting with a degree title or institution is a NEW entry
    4. Do NOT merge entries with same university or dates

    EXAMPLES (from this CV):
    [
    {{"institution": "Politechnika Śląska", "degree": "Doktorat", "field": "Mechanika Obliczeniowa", "period": "2018-2022"}},
    {{"institution": "Montanuniversität Leoben", "degree": "Staż doktorancki", "field": "energia odnawialna", "period": "2022"}},
    {{"institution": "Cranfield University", "degree": "Staż doktorancki", "field": "przepływy supersoniczne", "period": "2022"}},
    {{"institution": "Cranfield University", "degree": "Master of Science", "field": "Computational Fluid Dynamics", "period": "2017-2018"}},
    {{"institution": "Politechnika Śląska", "degree": "Magister Inżynier", "field": "Mechanika i Budowa Maszyn", "period": "2013-2017"}},
    {{"institution": "Politechnika Śląska", "degree": "Inżynier", "field": "Mechanika i Budowa Maszyn", "period": "2013-2017"}}
    ]

    CV TEXT:
    {section[:5500]}
    """

        # 3. LLM - większy limit dla 6 wpisów
        model = getattr(self, "model_name", getattr(self, "modelname", "qwen2.5:14b"))
        try:
            resp = ollama.chat(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                options={"temperature": 0.0, "num_predict": 2000},  # ← 2000 dla 6 wpisów
            )
            raw = resp["message"]["content"].strip()
            # print(f"📝 RAW LLM ({len(raw)} chars):")
            # print(raw[:600])
        except Exception as e:
            print(f"❌ LLM error: {e}")
            return []

        # 4. Clean JSON
        raw = raw.replace("``````", "").strip()
        l = raw.find("[")
        r = raw.rfind("]") + 1
        
        if l == -1 or r == 0 or r <= l:
            print("❌ No JSON array")
            return []

        json_text = raw[l:r]
        print(f"🔍 JSON array ({len(json_text)} chars):")
        print(json_text[:500])

        try:
            data = json.loads(json_text)
            if not isinstance(data, list):
                return []
        except Exception as e:
            print(f"❌ JSON parse error: {e}")
            print(f"Failed JSON: {json_text[:400]}")
            return []

        # 5. Normalizacja
        normalized = []
        for e in data:
            period = (
                e.get("period")
                or f"{e.get('start','')}-{e.get('end','')}".strip("- ")
                or e.get("dates", "")
            )

            institution = (e.get("institution") or e.get("university") or e.get("school") or "").strip()
            degree = (e.get("degree") or e.get("title") or e.get("level") or "").strip()
            field = (e.get("field") or e.get("major") or e.get("specialization") or "").strip()
            period = " ".join(period.split())

            if degree or institution:
                normalized.append({
                    "institution": institution,
                    "degree": degree,
                    "field": field,
                    "period": period,
                })

        print(f"✅ Extracted {len(normalized)} education entries:")
        # for i, e in enumerate(normalized, 1):
        #     print(f"  {i}. {e['degree']} {e['field']} @ {e['institution']} ({e['period']})")

        return normalized




    def _create_polish_prompt(self, cv_text, client_requirements, needs_translation=False, source_lang='polish'):
        """
        IMPROVED: Uses extracted work experience and education directly
        """
        
        # Pre-extract all work experience and education
        extracted_work_exp = self._extract_work_experience_details(cv_text)
        extracted_education = self._extract_education_details(cv_text)
        extracted_tech = self._extract_technologies_from_cv(cv_text)
        categorized_tech = self._categorize_technologies(extracted_tech)
        
        prompt = "Jesteś ekspertem HR specjalizującym się w analizie CV.\n\n"
        
        if needs_translation:
            prompt += f"WAŻNE: CV jest napisane po {self._get_language_name(source_lang, 'pl')}. "
            prompt += "Przeanalizuj je i wygeneruj raport PO POLSKU.\n\n"
        
        prompt += "TREŚĆ CV KANDYDATA:\n" + cv_text + "\n\n"
        prompt += "WYMAGANIA KLIENTA:\n" + client_requirements + "\n\n"
        
        # Extracted technologies
        prompt += "=" * 80 + "\n"
        prompt += "TECHNOLOGIE Z CV (UŻYJ WSZYSTKICH):\n"
        prompt += "=" * 80 + "\n\n"
        
        for category, techs in categorized_tech.items():
            if techs:
                prompt += f"{category}: {', '.join(techs)}\n"
        
        prompt += "\n" + "=" * 80 + "\n"
        
        # Extracted work experience
        prompt += "DOŚWIADCZENIE ZAWODOWE (UŻYJ WSZYSTKICH WPISÓW):\n"
        prompt += "=" * 80 + "\n\n"
        
        for idx, exp in enumerate(extracted_work_exp, 1):
            prompt += f"{idx}. {exp['period']} | {exp['company']} | {exp['position']}\n"
            if exp['description']:
                prompt += f"   Opis: {exp['description']}\n"
            if exp['technologies']:
                prompt += f"   Tech: {', '.join(exp['technologies'])}\n"
        
        prompt += "\n" + "=" * 80 + "\n"
        
        # Extracted education
        prompt += "EDUKACJA (UŻYJ WSZYSTKICH WPISÓW I DOKŁADNIE TEŻ DATY):\n"
        prompt += "=" * 80 + "\n\n"
        
        for idx, edu in enumerate(extracted_education):
            prompt += f""" {{
        "uczelnia": "{edu['institution']}",
        "stopien": "{edu['degree']}",
        "kierunek": "{edu['field']}",
        "okres": "{edu['period']}"
        }}{'' if idx == len(extracted_education) - 1 else ','}
        """
        
        prompt += "\n" + "=" * 80 + "\n\n"
        
        prompt += """Wygeneruj szczegółowy raport w formacie JSON PO POLSKU:

    🚨 KRYTYCZNE - OBOWIĄZKOWE:
    1. Liczba wpisów w "doswiadczenie_zawodowe" = liczba wpisów z powyższej listy
    2. Liczba wpisów w "wyksztalcenie" = liczba wpisów z powyższej listy
    3. KAŻDY wpis z listy MUSI być w JSON, NIE pomijaj ani nie łącz!
    4. Daty w edukacji - przepisz DOKŁADNIE jak w liście, bez zmian
    5. Technologie z każdego wpisu doświadczenia - WSZYSTKIE w "technologie"

    {
    "podstawowe_dane": {
        "imie_nazwisko": "Wyciągnij z CV",
        "email": "Email lub: nie podano",
        "telefon": "Telefon lub: nie podano"
    },
    "lokalizacja_i_dostepnosc": {
        "lokalizacja": "Dokładna lokalizacja TYLKO jeśli jasno podana, inaczej: nie podano",
        "preferencja_pracy_zdalnej": "Zdalna/Hybrydowa/Stacjonarna/nieokreślona",
        "dostepnosc": "Okres wypowiedzenia lub: nieokreślona"
    },
    "podsumowanie_profilu": "Krótka analiza 3-5 zdań kandydata",
    "doswiadczenie_zawodowe": ["""
        
        # Add each work experience entry
        for idx, exp in enumerate(extracted_work_exp):
            prompt += f"""    {{
        "okres": "{exp['period']}",
        "firma": "{exp['company']}",
        "stanowisko": "{exp['position']}",
        "kluczowe_osiagniecia": {json.dumps([exp['description']] if exp['description'] else [], ensure_ascii=False)},
        "obowiazki": [],
        "technologie": {json.dumps(exp['technologies'], ensure_ascii=False)}
        }}{'' if idx == len(extracted_work_exp) - 1 else ','}
    """
        
        prompt += """  ],
    "wyksztalcenie": ["""
        
        # Add each education entry
        for idx, edu in enumerate(extracted_education):
            prompt += f"""    {{
        "uczelnia": "{edu['uczelnia']}",
        "stopien": "{edu['stopien']}",
        "kierunek": "{edu['kierunek']}",
        "okres": "{edu['okres']}"
        }}{'' if idx == len(extracted_education) - 1 else ','}
    """
        
        prompt += """  ],
    "certyfikaty_i_kursy": [
        {
        "nazwa": "Nazwa certyfikatu",
        "typ": "certyfikat",
        "wystawca": "Organizacja",
        "data": "Rok"
        }
    ],
    "jezyki_obce": [
        {"jezyk": "Język ojczysty", "poziom": "Ojczysty"}
    ],
    "umiejetnosci": {
        "programowanie_skrypty": """ + json.dumps(categorized_tech.get('programming_scripting', []), ensure_ascii=False) + """,
        "frameworki_biblioteki": """ + json.dumps(categorized_tech.get('frameworks_libraries', []), ensure_ascii=False) + """,
        "infrastruktura_devops": """ + json.dumps(categorized_tech.get('infrastructure_devops', []), ensure_ascii=False) + """,
        "chmura": """ + json.dumps(categorized_tech.get('cloud', []), ensure_ascii=False) + """,
        "bazy_kolejki": """ + json.dumps(categorized_tech.get('databases_messaging', []), ensure_ascii=False) + """,
        "monitoring": """ + json.dumps(categorized_tech.get('monitoring', []), ensure_ascii=False) + """,
        "inne": """ + json.dumps(categorized_tech.get('other', []), ensure_ascii=False) + """
    },
    "podsumowanie_technologii": {
        "opis": "Podsumowanie głównych technologii",
        "glowne_technologie": """ + json.dumps(extracted_tech[:10], ensure_ascii=False) + """,
        "lata_doswiadczenia": "X lat"
    },
    "dopasowanie_do_wymagan": {
        "mocne_strony": ["Mocna strona 1", "Mocna strona 2"],
        "poziom_dopasowania": "wysoki/sredni/niski",
        "uzasadnienie": "Szczegółowe uzasadnienie",
        "rekomendacja": "TAK/NIE"
    }
    }
    """
        
        return prompt


    
    def _create_english_prompt(self, cv_text, client_requirements, needs_translation=False, source_lang='english'):
        """ENHANCED English prompt with pre-extracted technologies"""
        
        # STEP 1: Extract technologies first
        extracted_tech = self._extract_technologies_from_cv(cv_text)
        categorized_tech = self._categorize_technologies(extracted_tech)
        
        prompt = "You are an expert HR professional specializing in CV analysis.\n\n"
        
        if needs_translation:
            prompt += f"IMPORTANT: CV is written in {self._get_language_name(source_lang, 'en')}. "
            prompt += "Analyze it and generate a comprehensive report IN ENGLISH, translating all information.\n\n"
        
        prompt += "CV TEXT:\n" + cv_text + "\n\n"
        prompt += "CLIENT REQUIREMENTS:\n" + client_requirements + "\n\n"
        
        # CRITICAL: Provide extracted technologies
        prompt += "=" * 80 + "\n"
        prompt += "TECHNOLOGIES EXTRACTED FROM CV (USE ALL OF THEM!):\n"
        prompt += "=" * 80 + "\n\n"
        
        for category, techs in categorized_tech.items():
            if techs:
                prompt += f"{category}: {', '.join(techs)}\n"
        
        prompt += "\n" + "=" * 80 + "\n\n"
        
        prompt += "Generate a comprehensive report in JSON format IN ENGLISH:\n"
        prompt += '{\n'
        
        # Basic data
        prompt += ' "basic_data": {\n'
        prompt += ' "full_name": "Extract name and surname from CV",\n'
        prompt += ' "email": "Email or: not provided",\n'
        prompt += ' "phone": "Phone or: not provided"\n'
        prompt += ' },\n'
        
        # Location
        prompt += '"location_and_availability": {\n'
        prompt += '"location": "Exact location (city and country) ONLY if it is clearly stated in the CV. \n'
        prompt += 'If there are only generic hints (e.g. Poland, Remote, EU), set: unclear / not provided.",\n'
        prompt += '"remote_work_preference": "Remote / Hybrid / On-site or: not specified",\n'
        prompt += '"availability": "Notice period or: not specified"\n'
        prompt += '},\n'
        
        # Profile summary
        prompt += ' "profile_summary": "IMPORTANT: Write YOUR OWN analysis (3-5 sentences). Include: experience, competencies, match to requirements, recommendation.",\n'
        
        # Work experience
        prompt += ' "work_experience": [\n'
        prompt += ' {\n'
        prompt += ' "period": "YYYY - YYYY or YYYY - Present",\n'
        prompt += ' "company": "Company name",\n'
        prompt += ' "position": "Position title",\n'
        prompt += ' "key_achievements": ["List of achievements with specific numbers/results"],\n'
        prompt += ' "responsibilities": ["Optional - detailed responsibilities"],\n'
        prompt += ' "technologies": ["Technologies used in this period"]\n'
        prompt += ' }\n'
        prompt += ' ],\n'
        
        # Education
        prompt += ' "education": [\n'
        prompt += ' {\n'
        prompt += ' "institution": "University name",\n'
        prompt += ' "degree": "Bachelor/Master/PhD",\n'
        prompt += ' "field": "Field of study",\n'
        prompt += ' "period": "YYYY - YYYY"\n'
        prompt += ' }\n'
        prompt += ' ],\n'
        
        # Certifications
        prompt += ' "certifications_and_courses": [\n'
        prompt += ' {\n'
        prompt += ' "name": "Certification or course name",\n'
        prompt += ' "type": "certification or course",\n'
        prompt += ' "issuer": "Organization/Platform",\n'
        prompt += ' "date": "Year"\n'
        prompt += ' }\n'
        prompt += ' ],\n'
        
        # Languages
        prompt += ' "languages": [\n'
        prompt += ' {"language": "Language name", "level": "A1/A2/B1/B2/C1/C2/Native"}\n'
        prompt += ' ],\n'
        
        # Skills - USE EXTRACTED TECHNOLOGIES
        prompt += ' "skills": {\n'
        prompt += f' "programming_scripting": {json.dumps(categorized_tech["programming_scripting"])},\n'
        prompt += f' "frameworks_libraries": {json.dumps(categorized_tech["frameworks_libraries"])},\n'
        prompt += f' "infrastructure_devops": {json.dumps(categorized_tech["infrastructure_devops"])},\n'
        prompt += f' "cloud": {json.dumps(categorized_tech["cloud"])},\n'
        prompt += f' "databases_messaging": {json.dumps(categorized_tech["databases_messaging"])},\n'
        prompt += f' "monitoring": {json.dumps(categorized_tech["monitoring"])},\n'
        prompt += f' "other": {json.dumps(categorized_tech["other"])}\n'
        prompt += ' },\n'
        
        # Tech stack summary
        prompt += ' "tech_stack_summary": {\n'
        prompt += ' "description": "Brief summary of candidate main technologies",\n'
        prompt += ' "primary_technologies": ["Top 8-10 most important technologies from above list"],\n'
        prompt += ' "years_of_experience": "X years of IT experience"\n'
        prompt += ' },\n'
        
        # Matching
        prompt += ' "matching_to_requirements": {\n'
        prompt += ' "strengths": ["At least 3 strengths related to requirements"],\n'
        prompt += ' "match_level": "high/medium/low",\n'
        prompt += ' "justification": "Detailed justification with specific examples",\n'
        prompt += ' "recommendation": "YES - recommend for further process / NO - does not meet requirements"\n'
        prompt += ' }\n'
        prompt += '}\n\n'
        
        prompt += "CRITICAL INSTRUCTIONS:\n"
        prompt += "1. USE ALL technologies from the list above in the 'skills' section\n"
        prompt += "2. Extract ALL work experience information\n"
        prompt += "3. DO NOT ADD technologies that are not in the CV\n"
        prompt += "4. RETURN valid JSON with ALL fields filled\n\n"
        
        return prompt


    
    def translate_analysis_dict(analysis_dict, language="pl"):
        """Translate entire analysis dictionary to target language using LLM"""
        if language == "en":
            return analysis_dict  # No translation needed
        
        try:
            # Konwertuj dict do JSON string
            analysis_json = json.dumps(analysis_dict, ensure_ascii=False, indent=2)
            
            prompt = f"""Translate the following CV analysis from English to Polish.
    Keep all JSON structure intact. Keep names, dates, and technical terms unchanged.
    Only translate the text values.

    JSON to translate:
    {analysis_json}

    Respond ONLY with valid JSON, no other text."""
            
            response = ollama.generate(
                model="mistral",
                prompt=prompt,
                stream=False,
                temperature=0.1  # Low temperature for consistency
            )
            
            translated_text = response['response'].strip()
            
            # Spróbuj parsować JSON
            translated_dict = json.loads(translated_text)
            return translated_dict
            
        except Exception as e:
            print(f"Translation error: {e}")
            return analysis_dict  # Return original if translation fails


    def extract_key_highlights(self, client_requirements='', output_language='auto'):
        """
        Extract 1-2 SHORT bullets matching CV to requirements.
        ✅ FIXED: Checks SKILLS, LANGUAGES, CERTIFICATIONS, WORK EXPERIENCE, and EDUCATION
        """
        if not hasattr(self, '_current_cv_text') or not self._current_cv_text:
            print("⚠️ No CV text available for highlights")
            return []

        if not client_requirements or not client_requirements.strip():
            print("⚠️ No client requirements provided")
            return []

        cv_text = self._current_cv_text.lower()

        # Detect output language
        if output_language == 'auto':
            req_lower = client_requirements.lower()
            polish_count = sum(1 for word in ['znajomość', 'doświadczenie', 'umiejętność', 'lat']
                            if word in req_lower)
            english_count = sum(1 for word in ['knowledge', 'experience', 'skills', 'years']
                            if word in req_lower)
            output_language = 'polish' if polish_count > english_count else 'english'

        print(f"🌐 Highlight language: {output_language}")

        # ✅ CRITICAL: Define spoken languages (NOT programming languages)
        SPOKEN_LANGUAGES = {
            'polish': ['polski', 'polish', 'polacco'],
            'english': ['angielski', 'english', 'inglese'],
            'german': ['niemiecki', 'german', 'deutsch', 'tedesco'],
            'russian': ['rosyjski', 'russian', 'русский', 'russo'],
            'french': ['francuski', 'french', 'français', 'francese'],
            'spanish': ['hiszpański', 'spanish', 'español', 'spagnolo'],
            'italian': ['włoski', 'italian', 'italiano'],
            'chinese': ['chiński', 'chinese', '中文', 'cinese'],
            'japanese': ['japoński', 'japanese', '日本語', 'giapponese'],
            'ukrainian': ['ukraiński', 'ukrainian', 'українська'],
            'portuguese': ['portugalski', 'portuguese', 'português'],
            'dutch': ['holenderski', 'dutch', 'nederlands'],
            'swedish': ['szwedzki', 'swedish', 'svenska'],
            'norwegian': ['norweski', 'norwegian', 'norsk'],
            'danish': ['duński', 'danish', 'dansk'],
            'czech': ['czeski', 'czech', 'čeština'],
            'slovak': ['słowacki', 'slovak', 'slovenčina']
        }

        # ✅ NEW: Define education degree keywords
        EDUCATION_KEYWORDS = {
            'polish': {
                'bachelor': ['inżynier', 'licencjat', 'bachelor'],
                'master': ['magister', 'magister inżynier', 'master'],
                'phd': ['doktorat', 'phd', 'dr'],
                'engineer': ['inżynier', 'inżynierskie', 'engineer']
            },
            'english': {
                'bachelor': ['bachelor', 'bs', 'ba', 'beng'],
                'master': ['master', 'ms', 'ma', 'msc', 'meng'],
                'phd': ['phd', 'doctorate', 'doctoral'],
                'engineer': ['engineer', 'engineering', 'beng', 'meng']
            }
        }

        # Parse requirements
        req_lines = []
        for line in client_requirements.split('\n'):
            line = line.strip()
            if not line:
                continue
            # Remove bullets
            for marker in ['-', '•', '*', '◦', '○']:
                if line.startswith(marker):
                    line = line[len(marker):].strip()
                    break
            if line:
                req_lines.append(line)

        if not req_lines:
            print("⚠️ No valid requirements")
            return []

        print(f"📋 Requirements: {len(req_lines)}")
        for i, req in enumerate(req_lines, 1):
            print(f"  {i}. {req}")

        # Extract highlights
        import re
        highlights = []

        for req in req_lines:
            req_lower = req.lower()

            # ✅ STEP 1: Check if requirement is for SPOKEN LANGUAGE
            is_spoken_language = False
            spoken_lang_name = None
            spoken_lang_variants = []

            for lang_key, variants in SPOKEN_LANGUAGES.items():
                for variant in variants:
                    if variant in req_lower:
                        is_spoken_language = True
                        spoken_lang_name = lang_key.capitalize()
                        spoken_lang_variants = variants
                        break
                if is_spoken_language:
                    break

            # If SPOKEN LANGUAGE requirement detected
            if is_spoken_language:
                print(f"\n  🗣️ SPOKEN LANGUAGE detected: {spoken_lang_name}")
                print(f"     Variants to check: {spoken_lang_variants}")

                found_in_cv = False
                for variant in spoken_lang_variants:
                    if variant in cv_text:
                        found_in_cv = True
                        print(f"     ✅ Found '{variant}' in CV")
                        break

                if found_in_cv:
                    if output_language == 'polish':
                        highlight = f"{spoken_lang_name}: potwierdzony w sekcji języków"
                    else:
                        highlight = f"{spoken_lang_name}: confirmed in languages section"
                    highlights.append(highlight)
                    print(f"     ✅ HIGHLIGHT ADDED: {highlight}")
                else:
                    print(f"     ❌ {spoken_lang_name} NOT FOUND in CV - NO HIGHLIGHT")
                
                continue  # ✅ SKIP to next requirement

            # ✅ STEP 2: Check if requirement is for EDUCATION DEGREE
            is_education_req = False
            education_type = None
            education_keywords_to_check = []

            # Check both Polish and English keywords
            for lang in ['polish', 'english']:
                for degree_type, keywords in EDUCATION_KEYWORDS[lang].items():
                    for keyword in keywords:
                        if keyword in req_lower:
                            is_education_req = True
                            education_type = degree_type
                            education_keywords_to_check = EDUCATION_KEYWORDS['polish'][degree_type] + EDUCATION_KEYWORDS['english'][degree_type]
                            break
                    if is_education_req:
                        break
                if is_education_req:
                    break

            # If EDUCATION requirement detected
            if is_education_req:
                print(f"\n  🎓 EDUCATION detected: {education_type}")
                print(f"     Keywords to check: {education_keywords_to_check}")

                found_in_cv = False
                for keyword in education_keywords_to_check:
                    if keyword in cv_text:
                        found_in_cv = True
                        print(f"     ✅ Found '{keyword}' in CV")
                        break

                if found_in_cv:
                    if output_language == 'polish':
                        degree_names = {
                            'bachelor': 'Studia Inżynierskie',
                            'master': 'Studia Magisterskie',
                            'phd': 'Doktorat',
                            'engineer': 'Studia Inżynierskie'
                        }
                        degree_name = degree_names.get(education_type, education_type.capitalize())
                        highlight = f"{degree_name}: potwierdzone w CV"
                    else:
                        degree_names = {
                            'bachelor': "Bachelor's degree",
                            'master': "Master's degree",
                            'phd': 'PhD',
                            'engineer': "Engineering degree"
                        }
                        degree_name = degree_names.get(education_type, education_type.capitalize())
                        highlight = f"{degree_name}: confirmed in CV"
                    
                    highlights.append(highlight)
                    print(f"     ✅ HIGHLIGHT ADDED: {highlight}")
                else:
                    print(f"     ❌ {education_type.upper()} NOT FOUND in CV - NO HIGHLIGHT")
                
                continue  # ✅ SKIP to next requirement

            # ✅ STEP 3: If NOT spoken language or education, treat as TECHNOLOGY/SKILL
            # Extract technology from requirement
            tech_patterns = [
                r'znajomość\s+([A-Z][a-z]*(?:\+\+|#)?)',  # Polish: "Znajomość Python"
                r'knowledge\s+of\s+([A-Z][A-Za-z0-9+#\.]+)',  # English: "Knowledge of SQL"
                r'doświadczenie\s+(?:w|z)\s+([A-Z][A-Za-z0-9+#\.]+)',  # Polish: "Doświadczenie w Python"
                r'experience\s+(?:with|in)\s+([A-Z][A-Za-z0-9+#\.]+)',  # English: "Experience with C++"
                r'\b([A-Z][A-Za-z0-9+#\.]{2,})\b',  # Any capitalized tech word (Python, SQL, C++)
            ]

            tech_found = None
            for pattern in tech_patterns:
                match = re.search(pattern, req)
                if match:
                    tech_found = match.group(1)
                    break

            if not tech_found:
                # Fallback: take first capitalized word
                words = req.split()
                for word in words:
                    if len(word) > 2 and word[0].isupper():
                        tech_found = word.strip('.,;:')
                        break

            if not tech_found:
                print(f"     ❌ Could not extract technology from: {req}")
                continue

            # Check if this TECHNOLOGY exists in CV
            tech_lower = tech_found.lower()

            # IMPORTANT: Exclude spoken languages from tech matching
            is_actually_spoken_lang = False
            for variants in SPOKEN_LANGUAGES.values():
                if tech_lower in [v.lower() for v in variants]:
                    is_actually_spoken_lang = True
                    break

            if is_actually_spoken_lang:
                print(f"     ⚠️ {tech_found} is a SPOKEN LANGUAGE, not a technology - skipping")
                continue

            if tech_lower not in cv_text:
                print(f"     ❌ {tech_found} (tech) NOT FOUND in CV (skipping)")
                continue

            # Build highlight for TECHNOLOGY
            if output_language == 'polish':
                highlight = f"{tech_found}: doświadczenie potwierdzone w CV"
            else:
                highlight = f"{tech_found}: experience confirmed in CV"

            highlights.append(highlight)
            print(f"     ✅ {tech_found} (tech) FOUND in CV")

        print(f"\n📊 FINAL: {len(highlights)} highlights")
        for i, h in enumerate(highlights, 1):
            print(f"  {i}. {h}")

        return highlights




    def _get_language_name(self, lang_code, output_lang):
        """Get language name in specified language"""
        names = {
            'polish': {'pl': 'polsku', 'en': 'Polish'},
            'english': {'pl': 'angielsku', 'en': 'English'}
        }
        return names.get(lang_code, {}).get(output_lang, lang_code)
    

    def extract_certifications(self, cvtext: str) -> list:
        """
        Extract ALL certifications, courses, accomplishments, and awards from CV.
        Handles multiple section names and formats.
        """
        import json
        import re
        
        # 1. Wzorce dla różnych nazw sekcji
        patterns = [
            r'(?is)\b(Certifications?|CERTIFICATIONS?|Certificates?|CERTIFICATES?)\b.*?(?=(Education|EDUCATION|Languages|LANGUAGES|Skills|SKILLS|Projects|PROJECTS|Accomplishments|ACCOMPLISHMENTS|\Z))',
            r'(?is)\b(Courses?|COURSES?|Training|TRAINING|Kursy|KURSY)\b.*?(?=(Education|EDUCATION|Languages|LANGUAGES|Skills|SKILLS|Projects|PROJECTS|Nagrody|NAGRODY|\Z))',
            r'(?is)\b(Accomplishments?|ACCOMPLISHMENTS?|Achievements?|ACHIEVEMENTS?)\b.*?(?=(Education|EDUCATION|Languages|LANGUAGES|Skills|SKILLS|\Z))',
            r'(?is)\b(Awards?|AWARDS?|Nagrody|NAGRODY|Wyróżnienia|WYRÓŻNIENIA)\b.*?(?=(Education|EDUCATION|Languages|LANGUAGES|Skills|SKILLS|Publikacje|PUBLIKACJE|\Z))',
            r'(?is)\b(Additional Information|ADDITIONAL INFORMATION|Dodatkowe informacje|DODATKOWE INFORMACJE)\b.*?(?=(Education|EDUCATION|Languages|LANGUAGES|Skills|SKILLS|\Z))',
            r'(?is)\b(Kursy i Szkolenia|KURSY I SZKOLENIA)\b.*?(?=(Umiejętności|UMIEJĘTNOŚCI|Nagrody|NAGRODY|\Z))',
        ]
        
        sections = []
        for pattern in patterns:
            match = re.search(pattern, cvtext, re.DOTALL | re.IGNORECASE)
            if match:
                section = match.group(0).strip()
                if len(section) > 30:  # Minimum 30 znaków
                    sections.append(section)
                    print(f"📜 Found cert/course section: {match.group(1)} ({len(section)} chars)")
        
        # 2. Jeśli nie znaleziono sekcji, użyj ostatnich 4000 znaków
        if not sections:
            print("⚠️ Cert sections not found, using last 4000 chars")
            combined = cvtext[-4000:]
        else:
            combined = "\n\n".join(sections)
        
        print(f"📜 CERTIFICATIONS SECTION sent to LLM (first 500 chars):")
        print(combined[:500])
        print("=== END CERTIFICATIONS SECTION ===")
        
        prompt = f"""Extract ALL certifications, courses, training, awards, accomplishments, and achievements from this text.

    Return ONLY a valid JSON array, no markdown, no comments.

    Each entry:
    - "name": certification/course/award name (e.g. "SAS Certified Specialist", "Complete Python Developer", "TOEIC B2")
    - "type": "certification" OR "course" OR "award" OR "accomplishment"
    - "issuer": organization/platform (e.g. "SAS", "Udemy", "TOEIC", "University") - if not specified, use ""
    - "date": year or date (e.g. "2024", "2023-11", "19.02.2013") - if not specified, use ""
    - "details": additional info like duration, achievement description (optional)

    EXAMPLES:
    [
    {{"name": "SAS Certified Specialist: Base Programming Using SAS 9.4", "type": "certification", "issuer": "SAS", "date": "2024-08", "details": ""}},
    {{"name": "Complete Python Developer in 2023: Zero to Mastery", "type": "course", "issuer": "Udemy", "date": "2023-11", "details": "31h"}},
    {{"name": "English Certificate TOEIC - B2", "type": "certification", "issuer": "TOEIC", "date": "", "details": ""}},
    {{"name": "Trzecia najlepsza praca doktorska w dziedzinie metod obliczeniowych w Polsce", "type": "award", "issuer": "European Community on Computational Methods in Applied Sciences", "date": "2022", "details": ""}},
    {{"name": "Created fully functional ticket selling system for BEST organization", "type": "accomplishment", "issuer": "University", "date": "", "details": "Used in production for Prom party"}}
    ]

    CRITICAL RULES:
    1. Extract EVERYTHING from the text - certifications, courses, awards, accomplishments
    2. If issuer not clear, extract from context (e.g. "Udemy:", "SAS Certified" → issuer="Udemy"/"SAS")
    3. Keep name concise but descriptive
    4. Use "type" to categorize: certification/course/award/accomplishment

    If no certifications/courses/awards found, return [].

    TEXT:
    {combined[:3000]}
    """

        model = getattr(self, "model_name", getattr(self, "modelname", "qwen2.5:14b"))
        
        try:
            resp = ollama.chat(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                options={"temperature": 0.0, "num_predict": 1500},
            )
            raw = resp["message"]["content"].strip()
            print(f"📝 CERTIFICATIONS RAW LLM ({len(raw)} chars):")
            print(raw[:400])
            
            # Clean JSON
            raw = raw.replace("``````", "").strip()
            l = raw.find("[")
            r = raw.rfind("]") + 1
            
            if l == -1 or r == 0 or r <= l:
                print("❌ No JSON array for certifications")
                return []
            
            json_text = raw[l:r]
            print(f"🔍 Certifications JSON (first 300 chars): {json_text[:300]}")
            data = json.loads(json_text)
            
            if not isinstance(data, list):
                return []
            
            # Normalizacja
            normalized = []
            for cert in data:
                name = (cert.get("name") or cert.get("nazwa") or "").strip()
                cert_type = (cert.get("type") or cert.get("typ") or "certification").strip().lower()
                issuer = (cert.get("issuer") or cert.get("wystawca") or "").strip()
                date = (cert.get("date") or cert.get("data") or "").strip()
                details = (cert.get("details") or cert.get("szczegóły") or "").strip()
                
                if name:  # Minimum: nazwa certyfikatu/kursu
                    normalized.append({
                        "name": name,
                        "type": cert_type,
                        "issuer": issuer if issuer else "Not specified",
                        "date": date if date else "Not specified",
                        "details": details
                    })
            
            print(f"✅ Extracted {len(normalized)} certifications/courses/awards:")
            for c in normalized[:5]:  # Pokaż pierwsze 5
                print(f"  - {c['type'].upper()}: {c['name']} ({c['issuer']}, {c['date']})")
            
            return normalized
            
        except Exception as e:
            print(f"❌ Certifications extraction error: {e}")
            import traceback
            traceback.print_exc()
            return []


    def _extract_keywords_from_requirements(self, client_requirements):
        """Extract ONLY key technologies and experience years from client requirements"""
        if not client_requirements or len(client_requirements) < 5:
            return []
        
        keywords = []
        requirements_lower = client_requirements.lower()
        
        # 1. Extract ONLY technologies from knowledge base (no common words)
        for category, tech_list in self.TECH_KNOWLEDGE_BASE.items():
            for tech in tech_list:
                tech_lower = tech.lower()
                # Only add if it's a real technology (not common words)
                if tech_lower in requirements_lower and len(tech) > 2:
                    keywords.append(tech_lower)
        
        # 2. Extract experience years (numbers followed by 'years', 'lat', 'rok')
        import re
        # Match patterns like "5 lat", "3 years", "5+ lat"
        year_patterns = re.findall(r'(\d+)\+?\s*(?:lat|lata|year|years|rok|lata)', requirements_lower)
        keywords.extend(year_patterns)
        
        # 3. Extract seniority levels (ONLY if explicitly mentioned)
        seniority = ['senior', 'junior', 'mid-level', 'lead', 'principal', 'architect']
        for level in seniority:
            if level in requirements_lower:
                keywords.append(level)
        
        # Remove duplicates and filter out common words
        keywords = list(set(keywords))
        
        # FILTER OUT common words that shouldn't be highlighted
        excluded_words = [
            'experience', 'doświadczenie', 'znajomość', 'knowledge', 
            'umiejętność', 'praca', 'work', 'projekt', 'project',
            'aplikacja', 'application', 'system', 'develop', 'rozwój',
            'z', 'w', 'na', 'do', 'i', 'or', 'and', 'the', 'a', 'an'
        ]
        keywords = [kw for kw in keywords if kw not in excluded_words and len(kw) > 2]
        
        print(f"🎯 Filtered keywords (technologies + years): {keywords}")
        return keywords



    def _text_contains_keyword(self, text, keywords):
        """Check if text contains any of the keywords"""
        if not keywords or not text:
            return False
        
        text_lower = text.lower()
        return any(kw in text_lower for kw in keywords)


    def _write_text_with_underline(self, pdf, text, x, y, width, font_name, font_size, keywords, line_height=5, bold=False):
        """
        Write multi-line text with underlined keywords
        Returns final Y position
        """
        if not text or text in ["NA", "Nie podano w CV", "not provided"]:
            return y
        
        if not keywords:
            # No keywords - regular text
            pdf.set_xy(x, y)
            pdf.set_font(font_name, 'B' if bold else '', 9)
            pdf.multi_cell(width, line_height, text, align='L')
            return pdf.get_y()
        
        # Split into words
        words = str(text).split()
        current_line = []
        current_y = y
        
        pdf.set_font(font_name, 'B' if bold else '', 9)
        
        for word in words:
            test_line = ' '.join(current_line + [word])
            
            # Check if line would be too wide
            if pdf.get_string_width(test_line) > width and current_line:
                # Write current line
                pdf.set_xy(x, current_y)
                self._write_line_with_keywords(pdf, ' '.join(current_line), font_name, font_size, keywords, bold)
                current_y += line_height
                current_line = [word]
            else:
                current_line.append(word)
        
        # Write remaining words
        if current_line:
            pdf.set_xy(x, current_y)
            self._write_line_with_keywords(pdf, ' '.join(current_line), font_name, font_size, keywords, bold)
            current_y += line_height
        
        return current_y


    def _write_line_with_keywords(self, pdf, line, font_name, font_size, keywords, bold=False):
        """Write a single line with BOLD keywords (instead of underline)"""
        words = line.split()
        
        for i, word in enumerate(words):
            # Check if word contains keyword
            word_lower = word.lower().strip('.,;:!?()"\'')
            has_keyword = any(
                                kw.lower() == word_lower or kw.lower() in word_lower 
                                for kw in keywords
                            ) if keywords else False
            
            # Set style - BOLD for keywords
            if has_keyword:
                style = 'B'  # Bold for matching keywords
            elif bold:
                style = 'B'  # Bold if requested
            else:
                style = ''   # Normal text
            
            pdf.set_font(font_name, style, 9)
            pdf.write(font_size * 0.7, word + (' ' if i < len(words)-1 else ''))

    def extract_languages(self, cvtext: str) -> list:
        """Extract ALL languages with levels from CV using LLM."""
        import json
        import re
        
        # 1. Spróbuj znaleźć sekcję Languages/Języki
        patterns = [
            r'(?is)\b(Languages|LANGUAGES|Języki|JĘZYKI)\b.*?(?=(Skills|SKILLS|Umiejętności|UMIEJĘTNOŚCI|Education|EDUCATION|Wykształcenie|WYKSZTAŁCENIE|Accomplishments|ACCOMPLISHMENTS|\Z))',
            r'(?is)\b(Language|LANGUAGE|Język|JĘZYK)\b.*?(?=(Skills|SKILLS|Umiejętności|UMIEJĘTNOŚCI|Education|EDUCATION|Wykształcenie|WYKSZTAŁCENIE|\Z))',
        ]
        
        section = None
        for pattern in patterns:
            match = re.search(pattern, cvtext, re.DOTALL | re.IGNORECASE)
            if match:
                section = match.group(0).strip()
                print(f"🗣️ Found languages section ({len(section)} chars)")
                break
        
        # 2. Jeśli nie znaleziono sekcji, użyj ostatnich 3000 znaków (Languages są zwykle na końcu)
        if not section or len(section) < 50:
            print("⚠️ Languages section not found, using last 3000 chars of CV")
            section = cvtext[-3000:]
        
        print(f"🗣️ LANGUAGES SECTION sent to LLM (first 500 chars):")
        print(section[:500])
        print("=== END LANGUAGES SECTION ===")
        
        prompt = f"""Extract ALL LANGUAGES mentioned in this text with their proficiency levels.

    Return ONLY a valid JSON array, no markdown, no comments.

    Each entry:
    - "language": language name (e.g. "English", "Polish", "Spanish", "German")
    - "level": proficiency level (e.g. "C1", "B2", "A2", "Native", "First Language", "Advanced", "Intermediate")

    EXAMPLES:
    [
    {{"language": "Polish", "level": "First Language"}},
    {{"language": "English", "level": "C1"}},
    {{"language": "Spanish", "level": "B2"}},
    {{"language": "German", "level": "A2"}}
    ]

    If no languages found, return [].

    TEXT:
    {section[:4000]}
    """

        model = getattr(self, "model_name", getattr(self, "modelname", "qwen2.5:14b"))
        
        try:
            resp = ollama.chat(
                model=model,
                messages=[{"role": "user", "content": prompt}],
                options={"temperature": 0.0, "num_predict": 500},
            )
            raw = resp["message"]["content"].strip()
            print(f"📝 LANGUAGES RAW LLM ({len(raw)} chars):")
            print(raw[:300])
            
            # Clean JSON
            raw = raw.replace("``````", "").strip()
            l = raw.find("[")
            r = raw.rfind("]") + 1
            
            if l == -1 or r == 0 or r <= l:
                print("❌ No JSON array for languages")
                return []
            
            json_text = raw[l:r]
            print(f"🔍 Languages JSON: {json_text[:200]}")
            data = json.loads(json_text)
            
            if not isinstance(data, list):
                return []
            
            # Normalizacja
            normalized = []
            for lang in data:
                language = (lang.get("language") or lang.get("język") or "").strip()
                level = (lang.get("level") or lang.get("poziom") or "").strip()
                
                if language:  # Minimum: nazwa języka
                    normalized.append({
                        "language": language,
                        "level": level if level else "Not specified"
                    })
            
            print(f"✅ Extracted {len(normalized)} languages:")
            for l in normalized:
                print(f"  - {l['language']}: {l['level']}")
            
            return normalized
            
        except Exception as e:
            print(f"❌ Language extraction error: {e}")
            import traceback
            traceback.print_exc()
            return []


    def extract_raw_experience_block(self, cv_text):
        # Look for the "Doświadczenie zawodowe" (case-insensitive)
        import re
        start = re.search(r"doświadczenie zawodowe", cv_text, re.IGNORECASE)
        if not start:
            return ""
        start_idx = start.start()
        # Search for the next main header (e.g. "Wykształcenie", "Umiejętności", etc.)
        end = re.search(r"\n[A-ZĄĆĘŁŃÓŚŹŻ ]{5,}\n", cv_text[start_idx:], re.IGNORECASE)
        if end:
            end_idx = start_idx + end.start()
            return cv_text[start_idx:end_idx].strip()
        else:
            return cv_text[start_idx:].strip()
    def translate_skill_category(self, category_key, target_language='polish'):
        """
        Translate skill category keys to Polish or English.
        
        Args:
            category_key: English key like 'programming_scripting', 'frameworks_libraries'
            target_language: 'polish' or 'english'
        
        Returns:
            Translated category name
        """
        translations = {
            'programming_scripting': {'polish': 'Programowanie i skrypty', 'english': 'Programming & Scripting'},
            'frameworks_libraries': {'polish': 'Frameworki i biblioteki', 'english': 'Frameworks & Libraries'},
            'infrastructure_devops': {'polish': 'Infrastruktura i DevOps', 'english': 'Infrastructure & DevOps'},
            'cloud': {'polish': 'Chmura', 'english': 'Cloud'},
            'databases_messaging': {'polish': 'Bazy danych i kolejki', 'english': 'Databases & Messaging'},
            'mobile': {'polish': 'Mobile', 'english': 'Mobile'},
            'monitoring': {'polish': 'Monitoring', 'english': 'Monitoring'},
            'other': {'polish': 'Inne', 'english': 'Other'}
        }
        
        return translations.get(category_key, {}).get(target_language, category_key)  
    
    def translate_language_name(self, language_name, target_lang):
        """
        Dynamically translate language name using Ollama.
        
        Args:
            language_name: Language name in any form ("English", "angielski", etc.)
            target_lang: "polish" or "english" (REQUIRED)
        
        Returns:
            Translated language name
        """
        if not language_name or len(language_name.strip()) < 2:
            return language_name
        
        try:
            if target_lang == "polish":
                prompt = f"""Translate this language name to Polish. Return ONLY the Polish name, one word.

    Language: {language_name}

    Polish name:"""
            else:
                prompt = f"""Translate this language name to English. Return ONLY the English name, one word.

    Language: {language_name}

    English name:"""
            
            response = ollama.chat(
                model=self.model_name,
                messages=[{'role': 'user', 'content': prompt}],
                options={'temperature': 0.1, 'num_predict': 50}
            )
            
            translated = response['message']['content'].strip()
            
            # ✅ CLEAN UP - usuń prefiksy
            prefixes_to_remove = [
                "Polish name:", "English name:", "Polish:", "English:",
                "Translation:", "Nazwa polska:", "Angielska nazwa:"
            ]
            
            for prefix in prefixes_to_remove:
                if translated.lower().startswith(prefix.lower()):
                    translated = translated[len(prefix):].strip()
                    break
            
            # Usuń cudzysłowy
            translated = translated.strip('"\'')
            
            # Weź tylko pierwsze słowo (nazwa języka)
            translated = translated.split()[0].strip('.,!?')
            
            print(f"Language: {language_name} → {translated}")
            return translated
            
        except Exception as e:
            print(f"Language translation failed for {language_name}, keeping original")
            return language_name

        
    def generate_pdf_output(self, analysis, template_type='full', language=None, client_requirements=''):
        """Generate PDF with FPDF2 - Arsenal font - 2 pages layout"""
        import re
        from io import BytesIO
        filtered_analysis = self.apply_template_filters(analysis, template_type)
        template_mode = filtered_analysis.get('template_mode', template_type)

        if language is None:
            language = filtered_analysis.get('output_language', 'en')
        
        # Font paths
        arsenal_regular = r'C:\Users\Kamil Czyżewski\OneDrive - Integral Solutions sp. z o.o\Pulpit\Projects\HR_CV_Analyzer\arsenal\Arsenal-Regular.ttf'#"/app/arsenal/Arsenal-Regular.ttf" # 
        arsenal_bold =r'C:\Users\Kamil Czyżewski\OneDrive - Integral Solutions sp. z o.o\Pulpit\Projects\HR_CV_Analyzer\arsenal\Arsenal-Bold.ttf' #"/app/arsenal/Arsenal-Bold.ttf" #
        
        if template_mode == 'one_to_one':
            keywords = []
        else:
            keywords = self._extract_keywords_from_requirements(client_requirements)

        print(f"🔍 Extracted {len(keywords)} keywords for highlighting: {keywords[:10]}")
        print(f"📝 Client requirements: {client_requirements[:100]}...")

        def safe_text(text, default='N/A'):
            if text is None or text == '':
                return default
            return str(text)
        
        def add_section_header(pdf, x, y, title, max_width):
            """
            Rysuje nagłówek sekcji + cienką linię pod spodem.
            Zwraca nową pozycję Y tuż pod linią.
            """
            # ustaw pozycję i font
            pdf.set_xy(x, y)
            pdf.set_font("Arsenal", "B", 10)

            # tekst nagłówka (jedna linia / multi_cell jeśli długi)
            pdf.multi_cell(max_width, 5, title, align="L")

            # Y tuż pod tekstem
            y_after_text = pdf.get_y()

            # cienka linia pod nagłówkiem
            pdf.set_draw_color(76, 76, 76)
            pdf.set_line_width(0.3)
            pdf.line(x, y_after_text, x + max_width - 2, y_after_text)
            pdf.set_draw_color(0, 0, 0)

            # mały odstęp pod linią
            return y_after_text + 3
        def get_section_name(en_name):
            """Polish translation"""

            output_lang = filtered_analysis.get('output_language', 'english')
            translations = {
                'K E Y  H I G H L I G H T S':'P O D S U M O W A N I E  P R O F I L U',
                'E D U C A T I O N': 'W Y K S Z T A Ł C E N I E',
                'L A N G U A G E S': 'J Ę Z Y K I',
                'C E R T I F I C A T I O N S': 'C E R T Y F I K A T Y',
                'P R O F I L E  S U M M A R Y': 'R E K O M E N D A C J A',
                'S K I L L S': 'U M I E J Ę T N O Ś C I',
                'T E C H  S T A C K': 'T E C H N O L O G I E',
                'W O R K  E X P E R I E N C E': 'D O Ś W I A D C Z E N I E  Z A W O D O W E',
            }
            if output_lang == 'polish':
                return translations.get(en_name, en_name)
            return en_name

        # Get data
        work_exp_data = filtered_analysis.get("doswiadczenie_zawodowe") or filtered_analysis.get("work_experience", [])
        candidate_name = "CANDIDATE NAME"
        candidate_title = "Professional Title"
        
        if "podstawowe_dane" in filtered_analysis:
            candidate_name = safe_text(filtered_analysis["podstawowe_dane"].get('imie_nazwisko', 'CANDIDATE NAME')).upper()
        elif "personal_data" in filtered_analysis or "basic_data" in filtered_analysis:
            basic = filtered_analysis.get("personal_data") or filtered_analysis.get("basic_data")
            if basic:
                candidate_name = safe_text(basic.get('full_name') or basic.get('name') or 'CANDIDATE NAME').upper()
        
        if work_exp_data:
            candidate_title = safe_text(work_exp_data[0].get('stanowisko') or work_exp_data[0].get('position', 'Professional'))

        # Create PDF
        pdf = FPDF(orientation='P', unit='mm', format='A4')
        
        # Register fonts
        try:
            pdf.add_font('Arsenal', '', arsenal_regular)
            pdf.add_font('Arsenal', 'B', arsenal_bold)
            pdf.set_font('Arsenal', '', 10)
        except Exception as e:
            print(f"Font error: {e}")
            pdf.set_font('Helvetica', '', 10)
        
        pdf.set_margins(left=12.7, top=0, right=12.7)

        # ===== PAGE 1: HEADER + KEY HIGHLIGHTS =====
        if template_mode != 'one_to_one':
            pdf.add_page()
            # Blue header
            pdf.set_fill_color(50, 130, 180)
            pdf.rect(0, 0, 210, 40, 'F')
            
            # Logo
            
            try:
                pdf.image(logo_path, x=5, y=9, w=50)
            except Exception as e:
                print(f"Logo error: {e}")
            
            # Name - centered
            pdf.set_font('Arsenal', 'B', 24)
            pdf.set_text_color(255, 255, 255)
            pdf.set_xy(0, 12)
            pdf.cell(0, 8, candidate_name, align='C')
            
            # Title - centered
            pdf.set_font('Arsenal', '', 12)
            pdf.set_xy(0, 22)
            pdf.cell(0, 8, candidate_title, align='C')
            
            # Back to black
            pdf.set_text_color(0, 0, 0)
            pdf.set_y(50)
            pdf.set_x(12.7)

            # Generate/Display KEY HIGHLIGHTS
            highlights = filtered_analysis.get("key_highlights", [])
            profile_summary = filtered_analysis.get("podsumowanie_profilu") or filtered_analysis.get("profile_summary") or ""
            
            # Generate highlights if missing
            if not highlights:
                if profile_summary and profile_summary.strip():
                    if "•" in profile_summary:
                        highlights = [h.strip() for h in profile_summary.split("•") if h.strip()][:6]
                    else:
                        sentences = re.split(r'\.\s+(?=[A-Z])', profile_summary)
                        highlights = []
                        for s in sentences:
                            s = s.strip()
                            if len(s) > 10:
                                if not s.endswith('.'):
                                    s = s + '.'
                                highlights.append(s)
                            if len(highlights) >= 6:
                                break
            
            if highlights:
                pdf.set_font('Arsenal', 'B', 13)
                pdf.cell(0, 5, get_section_name('K E Y  H I G H L I G H T S'), ln=True)
                
                # Underline
                pdf.set_draw_color(76, 76, 76)
                pdf.set_line_width(0.3)
                y_before = pdf.get_y()
                pdf.line(12.7, y_before, 197.3, y_before)
                pdf.set_draw_color(0, 0, 0)
                
                pdf.set_y(y_before + 3)
                pdf.set_x(12.7)
                pdf.set_font('Arsenal', '', 13)
                
                for highlight in highlights:
                    highlight_text = safe_text(highlight).strip()
                    if highlight_text:
                        pdf.set_x(12.7)
                        current_y = self._write_text_with_underline(
                            pdf, f"• {highlight_text}", 12.7, pdf.get_y(),
                            185, 'Arsenal', 10, keywords, line_height=5
                        )
                        pdf.set_y(current_y)


        # ===== PAGE 2: TWO COLUMNS LAYOUT =====
        pdf.add_page()
        pdf.set_margins(left=12.7, top=0, right=12.7)

        # Blue header (repeat on page 2)
        pdf.set_fill_color(50, 130, 180)
        pdf.rect(0, 0, 210, 40, 'F')

        # Logo, Name, Title (same as page 1)
        try:
            pdf.image(logo_path, x=5, y=9, w=50)
        except Exception as e:
            print(f"Logo error: {e}")

        pdf.set_font('Arsenal', 'B', 24)
        pdf.set_text_color(255, 255, 255)
        pdf.set_xy(0, 12)
        pdf.cell(0, 8, candidate_name, align='C')

        pdf.set_font('Arsenal', '', 12)
        pdf.set_xy(0, 22)
        pdf.cell(0, 8, candidate_title, align='C')

        pdf.set_text_color(0, 0, 0)

        # STAŁE POZYCJE KOLUMN NA STRONIE 2
        page_y = 50.0
        col_left_x = 12.7
        col_right_x = 102.0      # ← Zmniejszone z 104.0
        col_left_w = 84.0        # ← Zmniejszone z 88.0  
        col_right_w = 84.0  
        
        def truncate_text(text, max_width, font_name, font_size, pdf):
            """Obcina tekst do max_width lub dodaje page break"""
            words = text.split()
            current_line = ""
            lines = []
            
            for word in words:
                test_line = f"{current_line} {word}".strip() if current_line else word
                pdf.set_font(font_name, '', font_size)
                pdf.set_xy(0, 0)  # test position
                w = pdf.get_string_width(test_line)
                if w > max_width - 8:  # 8mm margines bezpieczeństwa
                    if current_line:
                        lines.append(current_line)
                    current_line = word
                else:
                    current_line = test_line
            
            if current_line:
                lines.append(current_line)
            return lines
        # --- LEWA KOLUMNA: SKILLS, TECH, LANGUAGES, CERTS (bez page breaków) ---
        output_lang = filtered_analysis.get('output_language', 'english')
        skills_data = filtered_analysis.get("umiejetnosci") or filtered_analysis.get("skills")

        pdf.set_xy(col_left_x, page_y)
        y_left = page_y

        if skills_data:
            # nagłówek SKILLS dokładnie w page_y
            pdf.set_xy(col_left_x, page_y)
            pdf.set_font('Arsenal', 'B', 10)
            pdf.cell(col_left_w, 5, get_section_name('S K I L L S'), ln=True)
            y_line = pdf.get_y()
            pdf.set_draw_color(76, 76, 76)
            pdf.set_line_width(0.3)
            pdf.line(col_left_x, y_line, col_left_x + col_left_w - 2, y_line)
            pdf.set_draw_color(0, 0, 0)

            y_left = y_line + 3
            pdf.set_xy(col_left_x + 2, y_left)

            skill_cats = [
                ('programowanie_skrypty', 'programming_scripting'),
                ('frameworki_biblioteki', 'frameworks_libraries'),
                ('mobile', 'mobile'),
                ('infrastruktura_devops', 'infrastructure_devops'),
                ('chmura', 'cloud'),
                ('bazy_kolejki', 'databases_messaging'),
                ('monitoring', 'monitoring'),
                ('inne', 'other'),
            ]

            for pl_key, en_key in skill_cats:
                skills_list = skills_data.get(pl_key) or skills_data.get(en_key)
                if not skills_list:
                    continue

                category_name = self.translate_skill_category(
                    en_key,
                    'polish' if output_lang == 'polish' else 'english'
                )
                skills_str = ", ".join(safe_text(s) for s in skills_list)

                pdf.set_xy(col_left_x + 2, y_left)
                pdf.set_font('Arsenal', 'B', 10)
                pdf.multi_cell(col_left_w - 4, 4, f"{category_name}:", align='L')
                y_left = pdf.get_y()

                current_y = self._write_text_with_underline(
                    pdf, skills_str, col_left_x + 2, y_left,
                    col_left_w - 4, 'Arsenal', 9, keywords, line_height=3.8
                )
                y_left = current_y + 2
                pdf.set_xy(col_left_x + 2, y_left)

        # TECH STACK
        tech_summary = filtered_analysis.get("podsumowanie_technologii") or filtered_analysis.get("tech_stack_summary")
        if tech_summary:
            pdf.set_xy(col_left_x, y_left)
            y_left = add_section_header(pdf, col_left_x, y_left, get_section_name('T E C H  S T A C K'), col_left_w)
            pdf.set_xy(col_left_x + 2, y_left)

            if isinstance(tech_summary, dict):
                description = tech_summary.get('opis') or tech_summary.get('description', '')
                primary_tech = tech_summary.get('glownetechnologie') or tech_summary.get('primarytechnologies', [])
                years_exp = tech_summary.get('latadoswiadczenia') or tech_summary.get('yearsofexperience', '')

                if description:
                    description_short = safe_text(description)[:120]
                    pdf.set_font('Arsenal', '', 9)
                    pdf.multi_cell(col_left_w - 4, 3.5, description_short, align='L')
                    y_left = pdf.get_y() + 1

                if primary_tech:
                    tech_list = primary_tech if isinstance(primary_tech, list) else [primary_tech]
                    tech_list = [safe_text(t) for t in tech_list if safe_text(t) and safe_text(t) != 'N/A'][:6]
                    if tech_list:
                        main_label = "Główne:" if output_lang == 'polish' else "Main:"
                        pdf.set_xy(col_left_x + 2, y_left)
                        pdf.set_font('Arsenal', 'B', 7)
                        pdf.cell(15, 3, main_label, ln=False)
                        pdf.set_font('Arsenal', '', 9)
                        pdf.set_xy(col_left_x + 17, y_left)
                        pdf.multi_cell(col_left_w - 19, 3, ", ".join(tech_list), align='L')
                        y_left = pdf.get_y() + 1

                if years_exp:
                    years_exp_str = safe_text(years_exp)
                    m = re.search(r'(\d+)', years_exp_str)
                    if m:
                        years_num = m.group(1)
                        exp_label = "Dośw:" if output_lang == 'polish' else "Exp:"
                        pdf.set_xy(col_left_x + 2, y_left)
                        pdf.set_font('Arsenal', 'B', 9)
                        pdf.cell(0, 3,
                            f"{exp_label} {years_num}+ lat" if output_lang == 'polish'
                            else f"{exp_label} {years_num}+ yrs",
                            ln=True
                        )
                        y_left = pdf.get_y() + 1
            else:
                tech_text = safe_text(tech_summary)[:150]
                pdf.set_font('Arsenal', '', 9)
                pdf.multi_cell(col_left_w - 4, 3.5, tech_text, align='L')
                y_left = pdf.get_y() + 1

            y_left = pdf.get_y() + 2


        # LANGUAGES
        tech_end_y = y_left          # tu skończyły się technologie
        languages_start_y = tech_end_y + 4
        certs_start_y = languages_start_y + 22

        languages_data = filtered_analysis.get("języki_obce") or filtered_analysis.get("languages", [])
        if languages_data:
            pdf.set_xy(col_left_x, y_left)
            y_left = add_section_header(pdf, col_left_x, y_left, get_section_name("L A N G U A G E S"), col_left_w)
            pdf.set_xy(col_left_x + 2, y_left)

            max_rows = 4
            rows_used = 0
            for lang in languages_data:
                if rows_used >= max_rows:
                    break
                language = safe_text(lang.get("język") or lang.get("language"), "")[:30]
                level = safe_text(lang.get("poziom") or lang.get("level"), "")
                if language:
                    text = f"{language}: {level}" if level else language
                    pdf.set_font("Arsenal", "", 9)
                    pdf.multi_cell(col_left_w - 4, 4, text, align="L")
                    y_left = pdf.get_y() + 1
                    pdf.set_xy(col_left_x + 2, y_left)
                    rows_used += 1

            y_left = pdf.get_y() + 2

        # CERTIFICATIONS (ciąg dalszy flowa)
        certs_and_courses = (
            filtered_analysis.get("certyfikaty_i_kursy") or
            filtered_analysis.get("certifications_and_courses") or
            (filtered_analysis.get("certyfikaty", []) or []) +
            (filtered_analysis.get("certifications", []) or [])
        )
        if certs_and_courses:
            pdf.set_xy(col_left_x, y_left)
            y_left = add_section_header(
                pdf, col_left_x, y_left, get_section_name('C E R T I F I C A T I O N S'), col_left_w
            )
            pdf.set_xy(col_left_x + 2, y_left)

            max_rows = 3
            rows_used = 0
            for item in certs_and_courses:
                if rows_used >= max_rows:
                    break
                item_name = safe_text(item.get('nazwa') or item.get('name', ''))[:40]
                issuer = safe_text(item.get('wystawca') or item.get('issuer', ''))[:40]
                if item_name:
                    pdf.set_font('Arsenal', 'B', 9)
                    pdf.multi_cell(col_left_w - 4, 4, item_name, align='L')
                    y_left = pdf.get_y()
                    if issuer and issuer != 'Not specified':
                        pdf.set_font('Arsenal', '', 9)
                        pdf.multi_cell(col_left_w - 4, 3.5, issuer, align='L')
                        y_left = pdf.get_y()
                    y_left += 1
                    pdf.set_xy(col_left_x + 2, y_left)

            y_left = pdf.get_y() + 2


        # --- PRAWA KOLUMNA: start DOKŁADNIE w tym samym Y ---
        y_right = page_y
        pdf.set_xy(col_right_x, y_right)

        # PROFILE SUMMARY
        profile_summary = filtered_analysis.get('podsumowanie_profilu') or filtered_analysis.get('profile_summary')
        if template_mode != "onetoone" and profile_summary and profile_summary not in ("NA", "Nie podano w CV", "not provided", ""):
            y_right = add_section_header(
                pdf, col_right_x, y_right,
                get_section_name("P R O F I L E  S U M M A R Y"),
                col_right_w,
            )
            profile_text = safe_text(profile_summary).replace('•', '').strip()
            pdf.set_xy(col_right_x + 2, y_right)
            currenty = self._write_text_with_underline(
                pdf, profile_text, col_right_x + 2, pdf.get_y(),
                col_right_w - 4, "Arsenal", 9, keywords, line_height=4
            )
            y_right = currenty + 3

        # WORK EXPERIENCE (bez żadnego wpływu na lewą kolumnę)
        work_exp_data = filtered_analysis.get("doswiadczenie_zawodowe") or filtered_analysis.get("work_experience", [])
        if work_exp_data:
            pdf.set_xy(col_right_x, y_right)
            pdf.set_font('Arsenal', 'B', 10)
            pdf.cell(col_right_w, 5, get_section_name("W O R K  E X P E R I E N C E"), ln=True)
            y_line = pdf.get_y()
            pdf.set_draw_color(76, 76, 76)
            pdf.set_line_width(0.3)
            pdf.line(col_right_x, y_line, col_right_x + col_right_w - 2, y_line)
            pdf.set_draw_color(0, 0, 0)

            y_right = y_line + 3
            pdf.set_xy(col_right_x + 2, y_right)

            for idx, exp in enumerate(work_exp_data):
                if idx > 0:
                    y_right += 2  # ← Mniejszy odstęp

                period = safe_text(exp.get("okres") or exp.get("period"), "")
                company = safe_text(exp.get("firma") or exp.get("company"), "")[:30]  # ← OGRANICZENIE!
                position = safe_text(exp.get("stanowisko") or exp.get("position"), "")[:35]  # ← OGRANICZENIE!
                achievements = exp.get("kluczowe_osiagniecia") or exp.get("key_achievements") or []
                technologies = exp.get("technologie") or exp.get("technologies") or []

                # Period
                if period not in ("", "YYYY - YYYY", "Not specified", "NA"):
                    pdf.set_xy(col_right_x + 2, y_right)
                    pdf.set_font("Arsenal", "B", 8)
                    pdf.set_text_color(100, 100, 100)
                    pdf.cell(0, 4, period[:25], ln=True)  # ← OGRANICZENIE!
                    pdf.set_text_color(0, 0, 0)
                    y_right = pdf.get_y()

                # Position
                if position:
                    pdf.set_xy(col_right_x + 2, y_right)
                    pdf.set_font("Arsenal", "B", 9)
                    pdf.multi_cell(col_right_w - 6, 4, position, align="L")  # ← -6
                    y_right = pdf.get_y()

                # Company
                if company:
                    pdf.set_xy(col_right_x + 2, y_right)
                    pdf.set_font("Arsenal", "B", 8)
                    pdf.multi_cell(col_right_w - 6, 4, company, align="L")  # ← -6
                    y_right = pdf.get_y()

                # Achievements
                if achievements and isinstance(achievements, list):
                    pdf.set_font('Arsenal', '', 9)
                    for ach in achievements[:3]:  # ← MAX 3 na pozycję
                        bullet_text = safe_text(ach, '').strip()[:80]  # ← OGRANICZENIE!
                        if bullet_text:
                            pdf.set_xy(col_right_x + 4, y_right)
                            pdf.multi_cell(col_right_w - 8, 4, f"• {bullet_text}", align='L')  # ← -8
                            y_right = pdf.get_y()

                # Technologies
                if technologies:
                    tech_list = technologies if isinstance(technologies, list) else [technologies]
                    tech_list = [safe_text(t, "")[:15] for t in tech_list if safe_text(t, "")][:5]  # ← OGRANICZENIE!
                    tech_str = ", ".join(tech_list)
                    if tech_str:
                        pdf.set_xy(col_right_x + 4, y_right)
                        pdf.set_font("Arsenal", "", 8)  # ← Mniejszy font
                        pdf.set_text_color(80, 80, 80)
                        pdf.multi_cell(col_right_w - 8, 3.5, f"Tech: {tech_str}", align="L")  # ← -8
                        pdf.set_text_color(0, 0, 0)
                        y_right = pdf.get_y()
                y_right += 2

        # EDUCATION – prawa kolumna, bez page breaków
        education_data = filtered_analysis.get('wyksztalcenie') or filtered_analysis.get('education', [])
        if education_data:
            pdf.set_xy(col_right_x, y_right)
            y_right = add_section_header(
                pdf, col_right_x, y_right, get_section_name('E D U C A T I O N'), col_right_w
            )
            pdf.set_xy(col_right_x + 2, y_right)
            y_right += 1

            for edu in education_data:
                period = safe_text(edu.get('okres') or edu.get('period', ''))
                institution = safe_text(edu.get('uczelnia') or edu.get('institution', ''))
                degree = safe_text(edu.get('stopien') or edu.get('degree', ''))
                field = safe_text(edu.get('kierunek') or edu.get('field_of_study') or edu.get('field', ''))

                if period and period not in ['YYYY - YYYY', 'Not specified', '', 'N/A']:
                    pdf.set_xy(col_right_x + 2, y_right)
                    pdf.set_font('Arsenal', 'B', 8)
                    pdf.set_text_color(100, 100, 100)
                    pdf.cell(0, 4, period, ln=True)
                    pdf.set_text_color(0, 0, 0)
                    y_right = pdf.get_y()

                if field or degree:
                    text_fd = f"{field}, {degree}" if field and degree else (field or degree)
                    pdf.set_xy(col_right_x + 2, y_right)
                    pdf.set_font('Arsenal', 'B', 8)
                    pdf.multi_cell(col_right_w - 4, 4, text_fd, align='L')
                    y_right = pdf.get_y()

                if institution and institution not in ("Nazwa uczelni", "Not specified", "", "NA"):
                    pdf.set_xy(col_right_x + 2, y_right)
                    pdf.set_font('Arsenal', '', 9)
                    pdf.multi_cell(col_right_w - 4, 4, institution, align='L')
                    y_right = pdf.get_y()

                y_right += 2

        # Save to buffer
        buffer = BytesIO()
        pdf.output(buffer)
        buffer.seek(0)
        return buffer

    
    def generate_docx_output(self, analysis, template_type='full', language=None, client_requirements=''):
        """Generate DOCX with Arsenal font - all headers with underlines"""

        filtered_analysis = self.apply_template_filters(analysis, template_type)
        template_mode = filtered_analysis.get('template_mode', template_type)

        # one_to_one = brak boldowania z wymagań klienta
        if template_mode == 'one_to_one':
            keywords = []
        else:
            keywords = self._extract_keywords_from_requirements(client_requirements)


        if language is None:
            language = filtered_analysis.get('output_language', 'en')
        
        
        def safe_text(value, default=''):
            return str(value).strip() if value and str(value).strip() not in ['None', 'null', 'N/A'] else default
        
        def get_section_name(en_name):
            output_lang = filtered_analysis.get('output_language', 'english')
            translations = {
                'K E Y  H I G H L I G H T S': 'P O D S U M O W A N I E  P R O F I L U',
                'E D U C A T I O N': 'W Y K S Z T A Ł C E N I E',
                'L A N G U A G E S': 'J Ę Z Y K I',
                'C E R T I F I C A T I O N S': 'C E R T Y F I K A T Y',
                'P R O F I L E  S U M M A R Y': 'R E K O M E N D A C J A ',
                'S K I L L S': 'U M I E J Ę T N O Ś C I',
                'T E C H  S T A C K': 'T E C H N O L O G I E',
                'W O R K  E X P E R I E N C E': 'D O Ś W I A D C Z E N I E  Z A W O D O W E',
            }
            if output_lang == 'polish':
                return translations.get(en_name, en_name)
            return en_name
        
        def apply_arsenal_font(run, size=9, bold=False):
            """Apply Arsenal font styling"""
            run.font.name = 'Arsenal'
            run.font.size = Pt(size)
            if bold:
                run.font.bold = True
        
        def add_section_header_with_underline(cell, title):
            """Add header with underline using Arsenal font"""
            p = cell.add_paragraph()
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(2)
            
            run = p.add_run(title)
            apply_arsenal_font(run, size=10, bold=True)
            
            # Add underline
            pPr = p._element.get_or_add_pPr()
            pBdr = parse_xml(r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:bottom w:val="single" w:sz="12" w:space="1" w:color="4C4C4C"/></w:pBdr>')
            pPr.append(pBdr)
            
            return p
        
        def add_section_header_inline(cell, title):
            """Header z underline bez dodatkowych odstępów (do 1:1)"""
            p = cell.add_paragraph()
            # zero marginesów – kontrolujesz je sam przed wywołaniem
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)

            run = p.add_run(title)
            apply_arsenal_font(run, size=10, bold=True)

            pPr = p._element.get_or_add_pPr()
            pBdr = parse_xml(
                r'<w:pBdr xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                r'<w:bottom w:val="single" w:sz="12" w:space="1" w:color="4C4C4C"/>'
                r'</w:pBdr>'
            )
            pPr.append(pBdr)
            return p

        def create_full_width_header(doc, candidate_name, candidate_title):
            """Create full-width blue header"""
            header_table = doc.add_table(rows=1, cols=1)
            tbl = header_table._element
            tblPr = tbl.tblPr
            tr = header_table.rows[0]._tr
            trPr = tr.get_or_add_trPr()
            trHeight = OxmlElement('w:trHeight')
            trHeight.set(qn('w:val'), '900')  # ok. 1,8 cm; zwiększ do 1000–1200 jeśli chcesz więcej
            trHeight.set(qn('w:hRule'), 'atLeast')
            trPr.append(trHeight)
            # Szerokość i marginesy tabeli (jak było)
            tblW = OxmlElement('w:tblW')
            tblW.set(qn('w:w'), '6500')
            tblW.set(qn('w:type'), 'pct')
            tblPr.append(tblW)

            tblInd = OxmlElement('w:tblInd')
            tblInd.set(qn('w:w'), '-1440')
            tblInd.set(qn('w:type'), 'dxa')
            tblPr.append(tblInd)

            tblBorders = OxmlElement('w:tblBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'none')
                border.set(qn('w:sz'), '0')
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'auto')
                tblBorders.append(border)
            tblPr.append(tblBorders)

            tblCellMar = OxmlElement('w:tblCellMar')
            for margin_type in ['top', 'left', 'bottom', 'right']:
                margin = OxmlElement(f'w:{margin_type}')
                margin.set(qn('w:w'), '0')
                margin.set(qn('w:type'), 'dxa')
                tblCellMar.append(margin)
            tblPr.append(tblCellMar)

            # Niebieskie tło dla całego prostokąta
            header_cell = header_table.rows[0].cells[0]
            shading_elm = parse_xml(
                r'<w:shd {} w:fill="3282B4"/>'.format(
                    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                )
            )
            header_cell._element.get_or_add_tcPr().append(shading_elm)

            # Wewnątrz tej niebieskiej komórki robimy 2 kolumny: [LOGO][TEKST]
            inner_table = header_cell.add_table(rows=1, cols=2)
            inner_table.autofit = False

            logo_cell = inner_table.rows[0].cells[0]
            text_cell = inner_table.rows[0].cells[1]

            logo_cell.width = Inches(2.3)
            text_cell.width = Inches(5.2)

            # Usuwamy bordery wewnętrznej tabeli, żeby był czysty prostokąt
            for cell in (logo_cell, text_cell):
                tcPr = cell._element.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'none')
                    border.set(qn('w:sz'), '0')
                    border.set(qn('w:space'), '0')
                    border.set(qn('w:color'), 'auto')
                    tcBorders.append(border)
                tcPr.append(tcBorders)

            # LOGO – trochę wyżej
            logo_para = logo_cell.paragraphs[0]
            logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_para.paragraph_format.space_before = Pt(0)  # mniejsza / ujemna wartość = wyżej
            logo_para.paragraph_format.space_after = Pt(0)
            logo_para.paragraph_format.left_indent = Inches(0.3)

            logo_path = r'C:\Users\Kamil Czyżewski\OneDrive - Integral Solutions sp. z o.o\Pulpit\Projects\HR_CV_Analyzer\IS_New.png' #"/app/arsenal/Arsenal-Regular.ttf"# 
            try:
                logo_run = logo_para.add_run()
                logo_run.add_picture(logo_path, width=Inches(2.0))
            except Exception as e:
                print(f'❌ Logo error in DOCX: {e}')

            # TEKST (imię + stanowisko) – trochę wyżej
            text_para = text_cell.paragraphs[0]
            text_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            text_para.paragraph_format.space_before = Pt(2)
            text_para.paragraph_format.space_after = Pt(0)

            name_run = text_para.add_run(candidate_name)
            apply_arsenal_font(name_run, size=28, bold=True)
            name_run.font.color.rgb = RGBColor(255, 255, 255)

            title_para = text_cell.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.space_before = Pt(0)
            title_para.paragraph_format.space_after = Pt(2)  # było 10

            title_run = title_para.add_run(candidate_title)
            apply_arsenal_font(title_run, size=13, bold=False)
            title_run.font.color.rgb = RGBColor(255, 255, 255)

        work_exp_data = filtered_analysis.get("doswiadczenie_zawodowe") or filtered_analysis.get("work_experience", [])
        
        candidate_name = "CANDIDATE NAME"
        candidate_title = "Professional Title"
        
        if "podstawowe_dane" in filtered_analysis:
            candidate_name = safe_text(filtered_analysis["podstawowe_dane"].get('imie_nazwisko', 'CANDIDATE NAME')).upper()
        elif "personal_data" in filtered_analysis or "basic_data" in filtered_analysis:
                # ← DODAJ TEN WARUNEK DLA ANGIELSKIEJ WERSJI
            basic = filtered_analysis.get("personal_data") or filtered_analysis.get("basic_data")
            if basic:
                candidate_name = safe_text(basic.get('full_name') or basic.get('name') or 'CANDIDATE NAME').upper()
            
        if work_exp_data:
            candidate_title = safe_text(work_exp_data[0].get('stanowisko') or work_exp_data[0].get('position', 'Professional'))
        
        doc = Document()
        
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0)
            section.bottom_margin = Inches(0.3)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
        
        # ===== PAGE 1 =====
        if template_mode != 'one_to_one':
            create_full_width_header(doc, candidate_name, candidate_title)
            
            highlights = filtered_analysis.get("mocne_strony") or filtered_analysis.get("key_highlights", [])
            
            # Fallback tylko jeśli PUSTE
            if not highlights:
                profile_summary = filtered_analysis.get("podsumowanie_profilu") or filtered_analysis.get("profile_summary", "")
                if profile_summary:
                    # Najpierw spróbuj podzielić po • (jak w PDF)
                    highlights = [h.strip() for h in profile_summary.split('•') if h.strip()]
                    
                    # Jeśli za mało, podziel po zdaniach
                    if len(highlights) < 3:
                        sentences = re.split(r'\.\s+(?=[A-Z])', profile_summary)
                        highlights = []
                        for s in sentences:
                            s = s.strip()
                            if len(s) > 10:
                                if not s.endswith('.'):
                                    s += '.'
                                highlights.append(s)
                            if len(highlights) >= 6:
                                break

            profile_summary = filtered_analysis.get("podsumowanie_profilu") or filtered_analysis.get("profile_summary") or ""
            # PAGE 1 - KEY HIGHLIGHTS WITH UNDERLINE
            add_section_header_with_underline(doc, get_section_name('K E Y  H I G H L I G H T S'))

            if highlights:
                for highlight in highlights:
                    highlight_text = safe_text(highlight).strip()
                    if highlight_text:
                        # ✅ Dodaj bullet point
                        p = doc.add_paragraph(style='List Bullet')
                        p.paragraph_format.space_before = Pt(0)
                        p.paragraph_format.space_after = Pt(2)
                        p.paragraph_format.left_indent = Inches(0.15)  # Wcięcie dla bullet
                        p.paragraph_format.first_line_indent = Inches(-0.15)

                        # Dodaj tekst z bold keywords
                        words = str(highlight_text).split()
                        for i, word in enumerate(words):
                            word_lower = word.lower().strip('.,!?')
                            has_keyword = any(kw.lower() in word_lower for kw in keywords) if keywords else False
                            
                            run = p.add_run(word)
                            apply_arsenal_font(run, size=11, bold=has_keyword)
                            
                            if i < len(words) - 1:
                                run_space = p.add_run(' ')
                                apply_arsenal_font(run_space, size=11, bold=False)

            doc.add_page_break()
        else:
            # one_to_one – potrzebujemy tylko profile_summary do ewentualnego użycia niżej,
            # ale nie generujemy sekcji PROFILE SUMMARY ani pierwszej strony
            profile_summary = filtered_analysis.get("podsumowanie_profilu") or filtered_analysis.get("profile_summary") or ""

        
        # ===== PAGE 2: TWO COLUMNS =====
        create_full_width_header(doc, candidate_name, candidate_title)
        doc.add_paragraph()
        
        page2_table = doc.add_table(rows=1, cols=2)
        page2_table.autofit = False
        
        left_cell = page2_table.rows[0].cells[0]
        right_cell = page2_table.rows[0].cells[1]
        
        # Clear existing paragraphs
        for paragraph in list(left_cell.paragraphs):
            p = paragraph._element
            p.getparent().remove(p)
        for paragraph in list(right_cell.paragraphs):
            p = paragraph._element
            p.getparent().remove(p)
        
        # SKILLS WITH UNDERLINE
        if template_mode == 'one_to_one':
            # nagłówek „inline”, bez dodatkowych marginesów
            add_section_header_inline(
                left_cell,
                get_section_name('S K I L L S')
            )
        else:
            add_section_header_with_underline(
                left_cell,
                get_section_name('S K I L L S')
            )
        
        skills_data = filtered_analysis.get("umiejetnosci") or filtered_analysis.get("skills")
        if skills_data:
            # MAPA: (klucz_pl, klucz_en, LABEL_EN, LABEL_PL)
            skill_cats = [
                ('programowanie_skrypty', 'programming_scripting', 'Programming', 'Programowanie'),
                ('frameworki_biblioteki', 'frameworks_libraries', 'Frameworks', 'Frameworki'),
                ('mobile', 'mobile', 'Mobile', 'Mobile'),
                ('infrastruktura_devops', 'infrastructure_devops', 'Infrastructure', 'Infrastruktura'),
                ('chmura', 'cloud', 'Cloud', 'Chmura'),
                ('bazy_kolejki', 'databases_messaging', 'Data', 'Bazy danych'),
                ('monitoring', 'monitoring', 'Monitoring', 'Monitoring'),
                ('inne', 'other', 'Other', 'Inne'),
            ]
            
            # Pobierz język outputu
            output_lang = filtered_analysis.get('output_language', 'english')
            
            for pl_key, en_key, label_en, label_pl in skill_cats:
                skills_list = skills_data.get(pl_key) or skills_data.get(en_key)
                if skills_list:
                    # WYBIERZ LABEL według języka outputu
                    label = label_pl if output_lang == 'polish' else label_en
                    
                    skills_str = ', '.join([safe_text(s) for s in skills_list])

                    # Label - bold
                    p = left_cell.add_paragraph(f"{label}:")
                    p.paragraph_format.space_before = Pt(8)
                    p.paragraph_format.space_after = Pt(4)
                    for run in p.runs:
                        apply_arsenal_font(run, size=9, bold=True)
                    
                    # Technologies - with BOLD keywords
                    self._add_paragraph_with_bold_keywords(left_cell, skills_str, keywords, base_size=9)

        
        # TECH STACK WITH UNDERLINE
        p = left_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        add_section_header_with_underline(left_cell, get_section_name('T E C H  S T A C K'))
        
        tech_summary = filtered_analysis.get("podsumowanie_technologii") or filtered_analysis.get("tech_stack_summary")
        if tech_summary:
            description = tech_summary.get('opis') or tech_summary.get('description')
            if description:
                p = left_cell.add_paragraph(safe_text(description))
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    apply_arsenal_font(run, size=9, bold=False)
        
        # LANGUAGES WITH UNDERLINE
        p = left_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        add_section_header_with_underline(left_cell, get_section_name('L A N G U A G E S'))

        languages_data = filtered_analysis.get("jezyki_obce") or filtered_analysis.get("languages", [])
        if languages_data:
            for lang in languages_data:
                # NAPRAWIONE KLUCZE - sprawdź WSZYSTKIE możliwe nazwy
                language = safe_text(lang.get('język') or lang.get('jzyk') or lang.get('language') or lang.get('lang') or '')
                level = safe_text(lang.get('poziom') or lang.get('level') or '')
                
                # Pomiń puste
                if not language or language in ['', 'None', 'N/A', 'not provided']:
                    continue
                    
                # Format: "Angielski: C1" lub "English: Advanced"
                lang_text = f"{language}: {level}".strip()
                
                if lang_text and len(lang_text) > 3:
                    # Użyj _add_paragraph_with_bold_keywords jak w innych sekcjach
                    self._add_paragraph_with_bold_keywords(
                        left_cell, 
                        lang_text, 
                        keywords, 
                        base_size=9,
                        space_before=0,
                        space_after=0
                    )
        
        # CERTIFICATIONS WITH UNDERLINE
        p = left_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        add_section_header_with_underline(left_cell, get_section_name('C E R T I F I C A T I O N S'))
        
        certs_and_courses = (
            filtered_analysis.get("certyfikaty_i_kursy") or 
            filtered_analysis.get("certifications_and_courses") or
            (filtered_analysis.get("certyfikaty", []) or []) + (filtered_analysis.get("certifications", []) or [])
        )
        
        if certs_and_courses:
            for item in certs_and_courses:
                item_name = safe_text(item.get('nazwa') or item.get('name', ''))
                issuer = safe_text(item.get('wystawca') or item.get('issuer', ''))
                
                p = left_cell.add_paragraph(item_name)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    apply_arsenal_font(run, size=9, bold=True)
                
                p = left_cell.add_paragraph(issuer)
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                for run in p.runs:
                    apply_arsenal_font(run, size=9, bold=False)
        
        # ===== RIGHT COLUMN =====
        
        # PROFILE SUMMARY WITH UNDERLINE
        if template_mode != 'one_to_one' and profile_summary:
            add_section_header_with_underline(right_cell, get_section_name('P R O F I L E  S U M M A R Y'))
            profile_text = safe_text(profile_summary).replace('•', '').strip()

            self._add_paragraph_with_bold_keywords(right_cell, profile_text, keywords, base_size=9)
        
        # WORK EXPERIENCE WITH UNDERLINE
        if template_mode == 'one_to_one':
            # dokładnie tak samo jak SKILLS w 1:1
            add_section_header_inline(
                right_cell,
                get_section_name('W O R K  E X P E R I E N C E')
            )
        else:
            add_section_header_with_underline(
                right_cell,
                get_section_name('W O R K  E X P E R I E N C E')
            )

        #tu
        if work_exp_data:
            for idx, exp in enumerate(work_exp_data):
                period = safe_text(exp.get("okres") or exp.get("period"), '')
                company = safe_text(exp.get("firma") or exp.get("company"), '')
                position = safe_text(exp.get("stanowisko") or exp.get("position"), '')
                
                achievements = (
                    exp.get("kluczowe_osiagniecia") or 
                    exp.get("key_achievements") or 
                    exp.get("description") or 
                    []
                )
                
                if isinstance(achievements, str):
                    achievements = [achievements] if achievements.strip() else []
                
                technologies = exp.get("technologie") or exp.get("technologies") or []
                
                # Odstęp między jobami
                if idx > 0:
                    spacer = right_cell.add_paragraph()
                    spacer.paragraph_format.space_before = Pt(6)
                    spacer.paragraph_format.space_after = Pt(0)
                
                # 1. PERIOD
                if period and period not in ['', 'YYYY - YYYY', 'Not specified', 'N/A']:
                    p_period = right_cell.add_paragraph()
                    p_period.paragraph_format.space_before = Pt(0)
                    p_period.paragraph_format.space_after = Pt(0)
                    run_period = p_period.add_run(period)
                    apply_arsenal_font(run_period, size=9, bold=True)
                    run_period.font.color.rgb = RGBColor(100, 100, 100)
                
                # 2. POSITION
                if position and position not in ['', 'N/A', 'Not specified']:
                    p_position = right_cell.add_paragraph()
                    p_position.paragraph_format.space_before = Pt(0)
                    p_position.paragraph_format.space_after = Pt(0)
                    run_position = p_position.add_run(position)
                    apply_arsenal_font(run_position, size=9, bold=True)
                
                # 3. COMPANY
                if company and company not in ['', 'N/A', 'Not specified']:
                    p_company = right_cell.add_paragraph()
                    p_company.paragraph_format.space_before = Pt(0)
                    p_company.paragraph_format.space_after = Pt(4)
                    run_company = p_company.add_run(company)
                    apply_arsenal_font(run_company, size=9, bold=True)
                
                # 4. ACHIEVEMENTS - z BULLET POINTS
                if achievements and isinstance(achievements, list) and len(achievements) > 0:
                    for achievement in achievements:
                        achievement_text = safe_text(achievement, '').strip()
                        if achievement_text and len(achievement_text) > 3:
                            # ✅ Dodaj bullet point
                            p = right_cell.add_paragraph()
                            p.paragraph_format.space_before = Pt(0)
                            p.paragraph_format.space_after = Pt(2)
                            p.paragraph_format.left_indent = Inches(0.15)  # Wcięcie dla bullet
                            p.paragraph_format.first_line_indent = Inches(-0.15)  # Hanging indent
                            
                            # Dodaj symbol bullet
                            run_bullet = p.add_run('• ')
                            apply_arsenal_font(run_bullet, size=9, bold=False)
                            
                            # Dodaj tekst z bold keywords
                            words = str(achievement_text).split()
                            for i, word in enumerate(words):
                                word_lower = word.lower().strip('.,!?')
                                has_keyword = any(kw.lower() in word_lower for kw in keywords) if keywords else False
                                
                                run = p.add_run(word)
                                apply_arsenal_font(run, size=9, bold=has_keyword)
                                
                                if i < len(words) - 1:
                                    run_space = p.add_run(' ')
                                    apply_arsenal_font(run_space, size=9, bold=False)
                
                # 5. TECHNOLOGIES
                if technologies:
                    tech_list = technologies if isinstance(technologies, list) else [technologies]
                    tech_str = ', '.join([safe_text(t, '') for t in tech_list if safe_text(t, '')])
                    if tech_str:
                        p_tech = right_cell.add_paragraph()
                        p_tech.paragraph_format.space_before = Pt(0)
                        p_tech.paragraph_format.space_after = Pt(4)
                        run_tech = p_tech.add_run(f"Tech: {tech_str}")
                        apply_arsenal_font(run_tech, size=9, bold=False)
                        run_tech.font.color.rgb = RGBColor(80, 80, 80)
        
        # EDUCATION WITH UNDERLINE
        p = right_cell.add_paragraph()
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(0)
        add_section_header_with_underline(right_cell, get_section_name('E D U C A T I O N'))
        
        education_data = filtered_analysis.get('wyksztalcenie') or filtered_analysis.get('education', [])



        if education_data and len(education_data) > 0:
            
            for edu in education_data:
                institution = safe_text(edu.get('uczelnia') or edu.get('institution'), '')
                degree = safe_text(edu.get('stopien') or edu.get('degree'), '')
                field = safe_text(edu.get('kierunek') or edu.get('field_of_study') or edu.get('field'), '')
                period = safe_text(edu.get('okres') or edu.get('period'), '')
                
                # ✅ RENDERUJ TYLKO JEŚLI DANE ISTNIEJĄ (jak w PDF!)
                if period and period not in ['', 'None', 'N/A']:
                    p_period = right_cell.add_paragraph()
                    p_period.paragraph_format.space_before = Pt(0)
                    p_period.paragraph_format.space_after = Pt(0)
                    run_period = p_period.add_run(period)
                    apply_arsenal_font(run_period, size=9, bold=True)
                    run_period.font.color.rgb = RGBColor(100, 100, 100)

                # 2. DEGREE + FIELD (tytuł) - BOLD
                if degree or field:
                    degree_field = f"{degree}, {field}" if degree and field else (degree or field)
                    if degree_field and degree_field not in ['', 'None', 'N/A', ',']:
                        p_degree = right_cell.add_paragraph()
                        p_degree.paragraph_format.space_before = Pt(0)
                        p_degree.paragraph_format.space_after = Pt(0)
                        run_degree = p_degree.add_run(degree_field)
                        apply_arsenal_font(run_degree, size=9, bold=True)

                # 3. INSTITUTION (uczelnia) na końcu - zwykły font
                if institution and institution not in ['', 'None', 'N/A']:
                    p_inst = right_cell.add_paragraph()
                    p_inst.paragraph_format.space_before = Pt(0)
                    p_inst.paragraph_format.space_after = Pt(4)  # odstęp po całym wpisie
                    run_inst = p_inst.add_run(institution)
                    apply_arsenal_font(run_inst, size=9, bold=False)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer